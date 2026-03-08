import sys, os, json, markdown2, urllib.request, urllib.error, re
from pathlib import Path
from datetime import datetime
os.environ["QT_FONT_DPI"] = "115"

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTextEdit, QPushButton, QSplitter, QDialog, QLabel, QLineEdit,
    QFrame, QMessageBox, QListWidget, QListWidgetItem, QGridLayout,
    QScrollArea, QAbstractItemView, QTabWidget, QFileDialog,
    QColorDialog, QSpinBox, QCheckBox, QGraphicsScene, QGraphicsView,
    QGraphicsItem, QGraphicsRectItem, QGraphicsLineItem,
    QGraphicsTextItem, QGraphicsPathItem, QStackedWidget,
    QToolButton, QFontDialog, QMenu, QGraphicsEllipseItem, QGraphicsPixmapItem,
    QComboBox, QProgressDialog, QTreeWidget, QTreeWidgetItem
)
from PySide6.QtCore import (
    Qt, QUrl, QPoint, QPropertyAnimation, QEasingCurve, QTimer, QSize,
    QThread, QObject, Signal, QRectF, QPointF, QLineF, QRect
)
from PySide6.QtWidgets import QTextBrowser
from PySide6.QtGui import (
    QIcon, QFont, QPainter, QColor, QPen, QBrush, QImage, QPainterPath,
    QCursor, QPixmap, QAction, QSyntaxHighlighter, QTextCharFormat,
    QTextDocument, QTextCursor, QKeySequence, QShortcut, QTransform,
    QUndoStack, QUndoCommand
)

# ═══════════════════════════════════════════════════════════════════════════
# THEMES
# ═══════════════════════════════════════════════════════════════════════════
THEMES = [
    # 0 GitHub Dark
    dict(name="GitHub Dark",
         bg="#0d1117",    panel="#161b22",   border="#30363d",
         hover="#21262d", code_bg="#161b22", input_bg="#0d1117",
         text="#c9d1d9",  dim="#8b949e",     editor_fg="#e5e5e7",
         scroll="#3f3f46",accent="#8b5cf6",  accent2="#7c3aed",
         find_bg="#1a1a20", stats_sep="#3f3f46", line_num_bg="#0d1117",
         syn_h1="#79c0ff", syn_h2="#58a6ff", syn_bold="#e6edf3",
         syn_italic="#cba6f7", syn_code="#ff7b72", syn_fence="#8b949e",
         syn_img="#4ade80", syn_link="#58a6ff", syn_quote="#8b949e",
         syn_list="#f0883e", syn_hr="#30363d", syn_tag="#7ee787",
         syn_chk="#fbbf24", syn_alert="#d29922"),
    # 1 Dark #181818
    dict(name="Dark",
         bg="#181818",    panel="#1e1e1e",   border="#2d2d2d",
         hover="#252525", code_bg="#141414", input_bg="#181818",
         text="#d4d4d4",  dim="#808080",     editor_fg="#d4d4d4",
         scroll="#3a3a3a",accent="#8b5cf6",  accent2="#7c3aed",
         find_bg="#141414", stats_sep="#2d2d2d", line_num_bg="#181818",
         syn_h1="#9cdcfe", syn_h2="#569cd6", syn_bold="#d4d4d4",
         syn_italic="#c586c0", syn_code="#ce9178", syn_fence="#6a9955",
         syn_img="#4ec9b0", syn_link="#569cd6", syn_quote="#6a9955",
         syn_list="#dcdcaa", syn_hr="#2d2d2d", syn_tag="#4ec9b0",
         syn_chk="#fbbf24", syn_alert="#d7ba7d"),
    # 2 Solarized Dark
    dict(name="Solarized Dark",
         bg="#002b36",    panel="#073642",   border="#586e75",
         hover="#0a4455", code_bg="#073642", input_bg="#002b36",
         text="#839496",  dim="#586e75",     editor_fg="#93a1a1",
         scroll="#657b83",accent="#268bd2",  accent2="#2aa198",
         find_bg="#002b36", stats_sep="#586e75", line_num_bg="#002b36",
         syn_h1="#268bd2", syn_h2="#2aa198", syn_bold="#eee8d5",
         syn_italic="#6c71c4", syn_code="#dc322f", syn_fence="#657b83",
         syn_img="#859900", syn_link="#268bd2", syn_quote="#657b83",
         syn_list="#cb4b16", syn_hr="#586e75", syn_tag="#859900",
         syn_chk="#b58900", syn_alert="#cb4b16"),
]

_THEME_IDX = 0

def T() -> dict:
    return THEMES[_THEME_IDX]

def set_theme_index(idx: int):
    global _THEME_IDX
    _THEME_IDX = idx % len(THEMES)

# ═══════════════════════════════════════════════════════════════════════════
# RECENT FILES MANAGER
# ═══════════════════════════════════════════════════════════════════════════
RECENT_FILES_PATH = os.path.expanduser("~/.readme_builder_recent.json")
MAX_RECENT = 10

def load_recent_files():
    try:
        if os.path.exists(RECENT_FILES_PATH):
            with open(RECENT_FILES_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
    except:
        pass
    return []

def save_recent_files(files):
    try:
        with open(RECENT_FILES_PATH, "w", encoding="utf-8") as f:
            json.dump(files[:MAX_RECENT], f, ensure_ascii=False)
    except:
        pass

def add_recent_file(path):
    files = load_recent_files()
    path = os.path.abspath(path)
    if path in files:
        files.remove(path)
    files.insert(0, path)
    save_recent_files(files[:MAX_RECENT])

# ═══════════════════════════════════════════════════════════════════════════
# PREVIEW CSS TEMPLATE
# ═══════════════════════════════════════════════════════════════════════════
def _preview_css(t: dict) -> str:
    return f"""<style>
body{{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;
  font-size:16px;line-height:1.6;color:{t['text']};background:{t['bg']};
  padding:24px;max-width:1012px;margin:0 auto;}}
h1,h2,h3,h4,h5,h6{{font-weight:600;line-height:1.25;color:{t['text']};margin-top:24px;margin-bottom:16px;}}
h1{{font-size:2em;border-bottom:1px solid {t['border']};padding-bottom:.3em;}}
h2{{font-size:1.5em;border-bottom:1px solid {t['border']};padding-bottom:.3em;}}
p{{margin-top:0;margin-bottom:16px;}}
code{{font-family:ui-monospace,Consolas,monospace;background:rgba(110,118,129,.4);
  padding:.2em .4em;border-radius:6px;font-size:85%;}}
pre{{background:{t['code_bg']};border-radius:6px;padding:16px;overflow:auto;
  border:1px solid {t['border']};margin-bottom:16px;}}
pre code{{background:transparent;padding:0;font-size:100%;}}
img{{max-width:100%;border-radius:6px;}}
blockquote{{padding:0 1em;color:{t['dim']};border-left:.25em solid {t['border']};margin:0 0 16px;}}
table{{border-spacing:0;border-collapse:collapse;width:100%;margin-bottom:16px;}}
table th,table td{{padding:6px 13px;border:1px solid {t['border']};}}
table tr:nth-child(2n){{background:{t['code_bg']};}}
a{{color:{t['accent']};text-decoration:none;}}
a:hover{{text-decoration:underline;}}
hr{{height:.25em;padding:0;margin:24px 0;background:{t['border']};border:0;}}
details{{background:{t['code_bg']};border:1px solid {t['border']};border-radius:6px;padding:8px 16px;margin-bottom:16px;}}
summary{{cursor:pointer;font-weight:600;color:{t['accent']};}}
kbd{{background:{t['code_bg']};border:1px solid {t['border']};border-radius:4px;padding:2px 6px;font-size:85%;font-family:monospace;}}
.markdown-alert{{padding:8px 16px;margin-bottom:16px;border-left:.25em solid;border-radius:0 6px 6px 0;}}
.markdown-alert-note{{border-left-color:#2f81f7;background:rgba(47,129,247,.1);}}
.markdown-alert-warning{{border-left-color:#d29922;background:rgba(210,153,34,.1);}}
.markdown-alert-tip{{border-left-color:#3fb950;background:rgba(63,185,80,.1);}}
.markdown-alert-important{{border-left-color:#a371f7;background:rgba(163,113,247,.1);}}
.markdown-alert-caution{{border-left-color:#f85149;background:rgba(248,81,73,.1);}}
::-webkit-scrollbar{{width:8px;}}
::-webkit-scrollbar-track{{background:{t['bg']};}}
::-webkit-scrollbar-thumb{{background:{t['border']};border-radius:10px;}}
::-webkit-scrollbar-thumb:hover{{background:{t['scroll']};}}
</style>"""

# ═══════════════════════════════════════════════════════════════════════════
# OFFLINE TRANSLATION ENGINE
# ═══════════════════════════════════════════════════════════════════════════
TRANSLATION_DICT = {
    "ar": {
        # Common words
        "the": "ال", "a": "", "an": "",
        "project": "المشروع", "description": "الوصف", "installation": "التثبيت",
        "usage": "الاستخدام", "features": "المميزات", "contributing": "المساهمة",
        "license": "الرخصة", "requirements": "المتطلبات", "documentation": "التوثيق",
        "examples": "الأمثلة", "getting started": "البداية السريعة",
        "table of contents": "جدول المحتويات", "overview": "نظرة عامة",
        "introduction": "مقدمة", "setup": "الإعداد", "configuration": "الإعدادات",
        "author": "المؤلف", "contact": "التواصل", "acknowledgements": "الشكر والتقدير",
        "roadmap": "خارطة الطريق", "screenshots": "لقطات الشاشة",
        "built with": "تم البناء باستخدام", "tech stack": "التقنيات المستخدمة",
        "quick start": "البداية السريعة", "api reference": "مرجع API",
        "faq": "الأسئلة الشائعة", "changelog": "سجل التغييرات",
        "support": "الدعم", "security": "الأمان", "performance": "الأداء",
        "test": "اختبار", "testing": "الاختبار", "deployment": "النشر",
        "download": "تحميل", "install": "تثبيت", "run": "تشغيل",
        "build": "بناء", "compile": "تجميع", "execute": "تنفيذ",
        "update": "تحديث", "upgrade": "ترقية", "version": "الإصدار",
        "release": "الإصدار", "stable": "مستقر", "latest": "الأحدث",
        "deprecated": "مهجور", "experimental": "تجريبي", "beta": "بيتا",
        "note": "ملاحظة", "warning": "تحذير", "tip": "نصيحة",
        "important": "مهم", "caution": "تنبيه", "danger": "خطر",
        "required": "مطلوب", "optional": "اختياري", "default": "افتراضي",
        "example": "مثال", "output": "المخرجات", "input": "المدخلات",
        "parameter": "المعامل", "option": "الخيار", "flag": "العلامة",
        "command": "الأمر", "script": "السكريبت", "function": "الدالة",
        "class": "الكلاس", "method": "الميثود", "module": "الوحدة",
        "package": "الحزمة", "library": "المكتبة", "framework": "الإطار",
        "tool": "الأداة", "utility": "الأداة المساعدة", "plugin": "الإضافة",
        "extension": "الامتداد", "integration": "التكامل", "workflow": "سير العمل",
        "environment": "البيئة", "variable": "المتغير", "config": "الإعدادات",
        "database": "قاعدة البيانات", "server": "الخادم", "client": "العميل",
        "frontend": "الواجهة الأمامية", "backend": "الخادم الخلفي",
        "api": "API", "endpoint": "نقطة النهاية", "request": "الطلب",
        "response": "الاستجابة", "authentication": "المصادقة", "authorization": "التفويض",
        "error": "خطأ", "bug": "خلل", "fix": "إصلاح", "issue": "مشكلة",
        "feature": "ميزة", "improvement": "تحسين", "enhancement": "تحسين",
        "contributor": "مساهم", "maintainer": "المشرف", "owner": "المالك",
        "community": "المجتمع", "open source": "مفتوح المصدر",
        "fork": "نسخ", "clone": "استنساخ", "pull request": "طلب سحب",
        "merge": "دمج", "branch": "فرع", "commit": "إيداع",
        "repository": "المستودع", "remote": "البعيد", "origin": "الأصل",
        "status": "الحالة", "report": "تقرير", "feedback": "التغذية الراجعة",
        "contribute": "ساهم", "submit": "إرسال", "review": "مراجعة",
        "approve": "موافقة", "reject": "رفض", "close": "إغلاق",
        "open": "فتح", "create": "إنشاء", "delete": "حذف",
        "edit": "تعديل", "save": "حفظ", "load": "تحميل",
        "import": "استيراد", "export": "تصدير", "copy": "نسخ",
        "paste": "لصق", "cut": "قص", "undo": "تراجع", "redo": "إعادة",
        "search": "بحث", "find": "بحث", "replace": "استبدال",
        "filter": "فلتر", "sort": "ترتيب", "group": "تجميع",
        "format": "تنسيق", "style": "نمط", "theme": "المظهر",
        "font": "الخط", "color": "اللون", "size": "الحجم",
        "width": "العرض", "height": "الارتفاع", "margin": "الهامش",
        "padding": "الحشو", "border": "الحدود", "background": "الخلفية",
        "image": "صورة", "icon": "أيقونة", "logo": "الشعار",
        "screenshot": "لقطة شاشة", "preview": "معاينة", "demo": "عرض",
        "video": "فيديو", "animation": "رسوم متحركة", "gif": "صورة متحركة",
        "file": "ملف", "folder": "مجلد", "directory": "مسار",
        "path": "المسار", "url": "الرابط", "link": "رابط",
        "download": "تنزيل", "upload": "رفع", "sync": "مزامنة",
        "backup": "نسخة احتياطية", "restore": "استعادة", "reset": "إعادة تعيين",
        "initialize": "تهيئة", "start": "بدء", "stop": "إيقاف",
        "pause": "إيقاف مؤقت", "resume": "استئناف", "restart": "إعادة تشغيل",
        "enable": "تفعيل", "disable": "تعطيل", "toggle": "تبديل",
        "show": "عرض", "hide": "إخفاء", "display": "عرض",
        "list": "قائمة", "table": "جدول", "chart": "مخطط",
        "graph": "رسم بياني", "diagram": "مخطط", "flowchart": "مخطط انسيابي",
        "architecture": "البنية", "design": "التصميم", "pattern": "النمط",
        "structure": "الهيكل", "layout": "التخطيط", "component": "المكون",
        "system": "النظام", "platform": "المنصة", "service": "الخدمة",
        "application": "التطبيق", "app": "التطبيق", "program": "البرنامج",
        "software": "البرمجيات", "hardware": "الأجهزة", "device": "الجهاز",
        "mobile": "الجوال", "desktop": "سطح المكتب", "web": "الويب",
        "cloud": "السحابة", "network": "الشبكة", "internet": "الإنترنت",
        "local": "محلي", "global": "عالمي", "public": "عام",
        "private": "خاص", "protected": "محمي", "secure": "آمن",
        "encrypt": "تشفير", "decrypt": "فك التشفير", "hash": "تجزئة",
        "token": "رمز", "key": "مفتاح", "secret": "سري",
        "password": "كلمة المرور", "user": "المستخدم", "account": "الحساب",
        "profile": "الملف الشخصي", "settings": "الإعدادات", "preferences": "التفضيلات",
        "notification": "إشعار", "alert": "تنبيه", "message": "رسالة",
        "email": "البريد الإلكتروني", "phone": "الهاتف", "address": "العنوان",
        "name": "الاسم", "title": "العنوان", "label": "التسمية",
        "value": "القيمة", "result": "النتيجة", "output": "المخرج",
        "success": "نجاح", "failure": "فشل", "complete": "مكتمل",
        "progress": "التقدم", "loading": "تحميل", "processing": "معالجة",
        "pending": "معلق", "active": "نشط", "inactive": "غير نشط",
        "enabled": "مفعل", "disabled": "معطل", "available": "متاح",
        "unavailable": "غير متاح", "ready": "جاهز", "busy": "مشغول",
        "connected": "متصل", "disconnected": "منقطع", "online": "متصل",
        "offline": "غير متصل", "sync": "مزامنة", "async": "غير متزامن",
        "batch": "دفعة", "queue": "طابور", "cache": "ذاكرة التخزين المؤقت",
        "log": "سجل", "debug": "تصحيح", "trace": "تتبع",
        "monitor": "مراقبة", "analyze": "تحليل", "optimize": "تحسين",
        "deploy": "نشر", "release": "إصدار", "publish": "نشر",
        "distribute": "توزيع", "install": "تثبيت", "uninstall": "إزالة",
        "update": "تحديث", "migrate": "ترحيل", "rollback": "التراجع",
        "test": "اختبار", "debug": "تصحيح", "profile": "تحليل الأداء",
        "benchmark": "معيار الأداء", "coverage": "التغطية",
        "lint": "فحص الكود", "format": "تنسيق الكود",
        "compile": "تجميع", "build": "بناء", "bundle": "حزمة",
        "minify": "ضغط", "compress": "ضغط", "zip": "ضغط",
        "generate": "توليد", "scaffold": "هيكل", "template": "قالب",
        "boilerplate": "الكود الأساسي", "starter": "البداية",
        "starter kit": "مجموعة البداية", "seed": "البذرة",
        "mock": "محاكاة", "stub": "وهمي", "fake": "مزيف",
        "spy": "تجسس", "assert": "تأكيد", "expect": "توقع",
        "describe": "وصف", "context": "السياق", "scenario": "السيناريو",
        "given": "بفرض", "when": "عندما", "then": "عندها",
        "before": "قبل", "after": "بعد", "setup": "الإعداد",
        "teardown": "التفكيك", "fixture": "التحضير",
        "dependency": "تبعية", "injection": "حقن",
        "factory": "مصنع", "singleton": "وحيد", "observer": "مراقب",
        "decorator": "مزخرف", "adapter": "محول", "proxy": "وسيط",
        "interface": "واجهة", "abstract": "مجرد", "concrete": "ملموس",
        "inherit": "وراثة", "extend": "توسيع", "implement": "تنفيذ",
        "override": "تجاوز", "overload": "تحميل زائد",
        "callback": "استدعاء رجعي", "promise": "وعد",
        "async": "غير متزامن", "await": "انتظار",
        "stream": "دفق", "buffer": "مخزن مؤقت", "pipe": "أنبوب",
        "event": "حدث", "listener": "مستمع", "emit": "إطلاق",
        "trigger": "مشغل", "hook": "خطاف", "middleware": "وسيطة",
        "plugin": "إضافة", "extension": "امتداد", "addon": "وظيفة إضافية",
    },
    "fr": {
        "description": "Description", "installation": "Installation",
        "usage": "Utilisation", "features": "Fonctionnalités",
        "contributing": "Contribution", "license": "Licence",
        "requirements": "Prérequis", "documentation": "Documentation",
        "examples": "Exemples", "getting started": "Démarrage rapide",
        "table of contents": "Table des matières", "overview": "Aperçu",
        "introduction": "Introduction", "setup": "Configuration",
        "configuration": "Configuration", "author": "Auteur",
        "contact": "Contact", "acknowledgements": "Remerciements",
        "roadmap": "Feuille de route", "screenshots": "Captures d'écran",
        "built with": "Construit avec", "tech stack": "Stack technique",
        "quick start": "Démarrage rapide", "api reference": "Référence API",
        "faq": "Questions fréquentes", "changelog": "Journal des modifications",
        "support": "Support", "security": "Sécurité",
        "note": "Note", "warning": "Attention", "tip": "Conseil",
        "important": "Important", "caution": "Attention",
        "project": "Projet", "download": "Télécharger",
        "install": "Installer", "run": "Exécuter", "build": "Construire",
        "test": "Tester", "deploy": "Déployer", "update": "Mettre à jour",
    },
    "es": {
        "description": "Descripción", "installation": "Instalación",
        "usage": "Uso", "features": "Características",
        "contributing": "Contribución", "license": "Licencia",
        "requirements": "Requisitos", "documentation": "Documentación",
        "examples": "Ejemplos", "getting started": "Comenzando",
        "table of contents": "Tabla de contenidos", "overview": "Resumen",
        "introduction": "Introducción", "setup": "Configuración",
        "configuration": "Configuración", "author": "Autor",
        "contact": "Contacto", "acknowledgements": "Agradecimientos",
        "roadmap": "Hoja de ruta", "screenshots": "Capturas de pantalla",
        "built with": "Construido con", "tech stack": "Stack tecnológico",
        "quick start": "Inicio rápido", "api reference": "Referencia API",
        "faq": "Preguntas frecuentes", "changelog": "Registro de cambios",
        "support": "Soporte", "security": "Seguridad",
        "note": "Nota", "warning": "Advertencia", "tip": "Consejo",
        "important": "Importante", "caution": "Precaución",
        "project": "Proyecto", "download": "Descargar",
        "install": "Instalar", "run": "Ejecutar", "build": "Construir",
        "test": "Probar", "deploy": "Desplegar", "update": "Actualizar",
    },
    "de": {
        "description": "Beschreibung", "installation": "Installation",
        "usage": "Verwendung", "features": "Funktionen",
        "contributing": "Mitwirken", "license": "Lizenz",
        "requirements": "Anforderungen", "documentation": "Dokumentation",
        "examples": "Beispiele", "getting started": "Erste Schritte",
        "table of contents": "Inhaltsverzeichnis", "overview": "Überblick",
        "introduction": "Einführung", "setup": "Einrichtung",
        "configuration": "Konfiguration", "author": "Autor",
        "contact": "Kontakt", "acknowledgements": "Danksagungen",
        "roadmap": "Roadmap", "screenshots": "Screenshots",
        "built with": "Erstellt mit", "tech stack": "Tech-Stack",
        "quick start": "Schnellstart", "api reference": "API-Referenz",
        "faq": "Häufige Fragen", "changelog": "Änderungsprotokoll",
        "support": "Unterstützung", "security": "Sicherheit",
        "note": "Hinweis", "warning": "Warnung", "tip": "Tipp",
        "important": "Wichtig", "caution": "Vorsicht",
        "project": "Projekt", "download": "Herunterladen",
        "install": "Installieren", "run": "Ausführen", "build": "Erstellen",
        "test": "Testen", "deploy": "Bereitstellen", "update": "Aktualisieren",
    },
    "zh": {
        "description": "描述", "installation": "安装",
        "usage": "使用方法", "features": "功能特点",
        "contributing": "贡献指南", "license": "许可证",
        "requirements": "依赖要求", "documentation": "文档",
        "examples": "示例", "getting started": "快速开始",
        "table of contents": "目录", "overview": "概览",
        "introduction": "简介", "setup": "配置",
        "configuration": "配置说明", "author": "作者",
        "contact": "联系方式", "acknowledgements": "致谢",
        "roadmap": "路线图", "screenshots": "截图",
        "built with": "技术栈", "tech stack": "技术栈",
        "quick start": "快速开始", "api reference": "API 参考",
        "faq": "常见问题", "changelog": "更新日志",
        "support": "支持", "security": "安全",
        "note": "注意", "warning": "警告", "tip": "提示",
        "important": "重要", "caution": "注意",
        "project": "项目", "download": "下载",
        "install": "安装", "run": "运行", "build": "构建",
        "test": "测试", "deploy": "部署", "update": "更新",
    },
    "ja": {
        "description": "説明", "installation": "インストール",
        "usage": "使用方法", "features": "機能",
        "contributing": "コントリビューション", "license": "ライセンス",
        "requirements": "要件", "documentation": "ドキュメント",
        "examples": "例", "getting started": "はじめに",
        "table of contents": "目次", "overview": "概要",
        "introduction": "はじめに", "setup": "セットアップ",
        "configuration": "設定", "author": "著者",
        "contact": "連絡先", "acknowledgements": "謝辞",
        "roadmap": "ロードマップ", "screenshots": "スクリーンショット",
        "built with": "使用技術", "tech stack": "技術スタック",
        "quick start": "クイックスタート", "api reference": "APIリファレンス",
        "faq": "よくある質問", "changelog": "変更履歴",
        "support": "サポート", "security": "セキュリティ",
        "note": "注意", "warning": "警告", "tip": "ヒント",
        "important": "重要", "caution": "注意",
        "project": "プロジェクト", "download": "ダウンロード",
        "install": "インストール", "run": "実行", "build": "ビルド",
        "test": "テスト", "deploy": "デプロイ", "update": "更新",
    },
    "ru": {
        "description": "Описание", "installation": "Установка",
        "usage": "Использование", "features": "Возможности",
        "contributing": "Участие", "license": "Лицензия",
        "requirements": "Требования", "documentation": "Документация",
        "examples": "Примеры", "getting started": "Начало работы",
        "table of contents": "Содержание", "overview": "Обзор",
        "introduction": "Введение", "setup": "Настройка",
        "configuration": "Конфигурация", "author": "Автор",
        "contact": "Контакты", "acknowledgements": "Благодарности",
        "roadmap": "Дорожная карта", "screenshots": "Скриншоты",
        "built with": "Создано с помощью", "tech stack": "Стек технологий",
        "quick start": "Быстрый старт", "api reference": "Справочник API",
        "faq": "Часто задаваемые вопросы", "changelog": "Журнал изменений",
        "support": "Поддержка", "security": "Безопасность",
        "note": "Примечание", "warning": "Предупреждение", "tip": "Совет",
        "important": "Важно", "caution": "Осторожно",
        "project": "Проект", "download": "Скачать",
        "install": "Установить", "run": "Запустить", "build": "Собрать",
        "test": "Тестировать", "deploy": "Развернуть", "update": "Обновить",
    },
}

LANG_NAMES = {
    "ar": "Arabic",
    "fr": "French",
    "es": "Spanish",
    "de": "German",
    "zh": "Chinese",
    "ja": "Japanese",
    "ru": "Russian",
}

def translate_markdown_offline(text: str, target_lang: str) -> str:
    """Fast offline translation of markdown headings and common section titles."""
    if target_lang not in TRANSLATION_DICT:
        return text
    
    d = TRANSLATION_DICT[target_lang]
    is_rtl = target_lang == "ar"
    lines = text.split("\n")
    result = []
    
    for line in lines:
        # Only translate heading lines and short labels, keep code/links intact
        stripped = line.lstrip("#").strip()
        lower = stripped.lower()
        
        # Check if it's a heading
        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            hashes = heading_match.group(1)
            content = heading_match.group(2)
            content_lower = content.lower()
            # Remove emojis for translation matching
            clean_content = re.sub(r'[^\w\s]', '', content_lower).strip()
            
            translated = None
            # Try multi-word matches first (longer keys first)
            for key in sorted(d.keys(), key=len, reverse=True):
                if key in clean_content:
                    translated = d[key]
                    break
            
            if translated:
                if is_rtl:
                    result.append(f"{hashes} {translated}")
                else:
                    result.append(f"{hashes} {translated}")
            else:
                result.append(line)
        else:
            result.append(line)
    
    translated_text = "\n".join(result)
    
    # Add RTL direction hint for Arabic
    if is_rtl:
        translated_text = "<!-- dir: rtl -->\n" + translated_text
    
    return translated_text


# ═══════════════════════════════════════════════════════════════════════════
# i18n
# ═══════════════════════════════════════════════════════════════════════════
SNIPPETS_FILE = os.path.expanduser("~/.readme_builder_snippets.json")
CURR_LANG = "en"
LANGS = {
    "en": {
        "app_title":"README Builder Pro","smart":"Smart","draw":"Draw",
        "score":"Score","import":"Import","snippet":"Snippet",
        "emojis":"Emojis","toc":"TOC","quick":"Quick","wizard":"Wizard",
        "save_exit":"Save & Exit","save":"Saved!","back":"Back",
        "Python Library":"Python Library","Web Application":"Web Application",
        "Machine Learning":"Machine Learning","Mobile App":"Mobile App",
        "CLI Tool":"CLI Tool","Game":"Game","Docker / DevOps":"Docker / DevOps",
        "Data Analysis":"Data Analysis","REST API":"REST API","UI Kit":"UI Kit",
        "Security Tool":"Security Tool","Documentation":"Documentation",
        "Organization":"Organization","GitHub Profile":"GitHub Profile",
        "VS Code Extension":"VS Code Extension","Discord Bot":"Discord Bot",
        "Cloud / SaaS":"Cloud / SaaS","Tutorial / Course":"Tutorial / Course",
        "Research Project":"Research Project","Blockchain":"Blockchain",
        "Desktop App":"Desktop App","Chrome Extension":"Chrome Extension",
        "NPM Package":"NPM Package","Arduino":"Arduino",
    },
    "ar": {
        "app_title":"منشئ README المحترف","smart":"ذكي","draw":"رسم",
        "score":"تقييم","import":"استيراد","snippet":"اختصار",
        "emojis":"إيموجي","toc":"فهرس","quick":"سريع","wizard":"مساعد",
        "save_exit":"حفظ وخروج","save":"تم الحفظ!","back":"رجوع",
    }
}
def tr(key): return LANGS[CURR_LANG].get(key, key)


QUICK_TEMPLATES = {
    tr("Python Library"): ("# {name}\n\n[![PyPI](https://img.shields.io/pypi/v/{slug}?style=flat-square)](#)\n[![Python](https://img.shields.io/badge/Python-3.8+-blue?style=flat-square&logo=python)](#)\n[![License](https://img.shields.io/badge/License-MIT-green?style=flat-square)](#)\n\n> Short description of what this library does.\n\n## Install\n```bash\npip install {slug}\n```\n\n## Quick Start\n```python\nimport {slug}\n\nresult = {slug}.do_something()\nprint(result)\n```\n\n## API Reference\n| Function | Description | Returns |\n|----------|-------------|--------|\n| `func()` | Does X | `str` |\n\n## Contributing\nSee [CONTRIBUTING.md](CONTRIBUTING.md)\n\n## License\nMIT © [Author](https://github.com/author)\n"),
    tr("Web Application"): ("# {name}\n\n[![Demo](https://img.shields.io/badge/Live-Demo-blue?style=flat-square)](#)\n[![React](https://img.shields.io/badge/React-18+-61DAFB?style=flat-square&logo=react)](#)\n\n<div align='center'>\n\n![Preview](screenshots/preview.png)\n\n</div>\n\n## Features\n- [ ] Feature one\n- [ ] Feature two\n- [ ] Feature three\n\n## Tech Stack\n| Layer | Tech |\n|-------|------|\n| Frontend | React / Vue |\n| Backend | Node / FastAPI |\n| Database | PostgreSQL |\n\n## Getting Started\n```bash\ngit clone https://github.com/user/{slug}.git\ncd {slug}\nnpm install && npm run dev\n```\n\n## Screenshots\n| Home | Dashboard |\n|------|----------|\n| ![](ss/home.png) | ![](ss/dash.png) |\n\n## License\nMIT\n"),
    tr("Machine Learning"): ("# {name}\n\n[![Python](https://img.shields.io/badge/Python-3.9+-blue?style=flat-square&logo=python)](#)\n[![PyTorch](https://img.shields.io/badge/PyTorch-2.0-EE4C2C?style=flat-square&logo=pytorch)](#)\n[![Paper](https://img.shields.io/badge/Paper-arXiv-red?style=flat-square)](#)\n\n## Abstract\n> Write your abstract here.\n\n## Results\n| Model | Accuracy | F1 |\n|-------|----------|----|  \n| Ours | **95.3%** | **0.94** |\n| Baseline | 92.1% | 0.91 |\n\n## Training\n```bash\npip install -r requirements.txt\npython train.py --config config.yaml\n```\n\n## Inference\n```python\nfrom model import Model\nm = Model.from_pretrained('checkpoint/')\noutput = m.predict(input_data)\n```\n\n## Citation\n```bibtex\n@article{name2024,\n  title={{name}},\n  author={Author},\n  year={2024}\n}\n```\n"),
    tr("Mobile App"): ("# {name}\n\n[![iOS](https://img.shields.io/badge/iOS-14+-000?style=flat-square&logo=apple)](#)\n[![Android](https://img.shields.io/badge/Android-9+-3DDC84?style=flat-square&logo=android)](#)\n\n<div align='center'>\n\n![App Preview](assets/mockup.png)\n\n</div>\n\n## Download\n[![App Store](https://img.shields.io/badge/App_Store-Download-blue?style=flat-square&logo=apple)](#)\n[![Google Play](https://img.shields.io/badge/Google_Play-Download-green?style=flat-square&logo=google-play)](#)\n\n## Features\n- Secure auth\n- Dark/Light mode\n- Offline support\n- Push notifications\n\n## Build\n```bash\nnpm install\nnpx react-native run-ios\nnpx react-native run-android\n```\n\n## License\nMIT\n"),
    tr("CLI Tool"): ("# {name}\n\n[![npm](https://img.shields.io/npm/v/{slug}?style=flat-square)](#)\n[![Downloads](https://img.shields.io/npm/dm/{slug}?style=flat-square)](#)\n\n> A powerful CLI tool for ...\n\n## Install\n```bash\nnpm install -g {slug}\n```\n\n## Usage\n```bash\n{slug} [command] [options]\n\nCommands:\n  init      Initialize project\n  build     Build the project\n  deploy    Deploy to server\n\nOptions:\n  -h, --help     Show help\n  -v, --version  Show version\n```\n\n## Config\n```yaml\n# .{slug}rc.yml\noutput: dist/\nverbose: false\n```\n\n## License\nMIT\n"),
    tr("REST API"): ("# {name} API\n\n[![FastAPI](https://img.shields.io/badge/FastAPI-0.100+-009688?style=flat-square&logo=fastapi)](#)\n[![Swagger](https://img.shields.io/badge/Docs-Swagger-85EA2D?style=flat-square)](#)\n\n## Base URL\n```\nhttps://api.example.com/v1\n```\n\n## Authentication\n```bash\ncurl -H 'Authorization: Bearer TOKEN' https://api.example.com/v1/items\n```\n\n## Endpoints\n| Method | Endpoint | Description |\n|--------|----------|-------------|\n| `GET` | `/items` | List all |\n| `POST` | `/items` | Create |\n| `GET` | `/items/{id}` | Get one |\n| `PUT` | `/items/{id}` | Update |\n| `DELETE` | `/items/{id}` | Delete |\n\n## Request Body\n```json\n{{\n  \"name\": \"New Item\",\n  \"description\": \"...\"\n}}\n```\n\n## Run\n```bash\npip install fastapi uvicorn\nuvicorn main:app --reload\n```\n\n## License\nMIT\n"),
    tr("Docker / DevOps"): ("# {name}\n\n[![Docker](https://img.shields.io/badge/Docker-Ready-2496ED?style=flat-square&logo=docker)](#)\n[![CI](https://img.shields.io/github/actions/workflow/status/user/{slug}/ci.yml?style=flat-square)](#)\n\n## Quick Start\n```bash\ndocker pull user/{slug}\ndocker run -d -p 8080:8080 user/{slug}\n```\n\n## Docker Compose\n```yaml\nversion: '3.8'\nservices:\n  app:\n    image: user/{slug}\n    ports:\n      - '8080:8080'\n  db:\n    image: postgres:15\n```\n\n## Env Variables\n| Variable | Default | Description |\n|----------|---------|-------------|\n| `PORT` | `8080` | Server port |\n| `DB_URL` | - | Database URL |\n\n## License\nMIT\n"),
    tr("Data Analysis"): ("# {name}\n\n[![Python](https://img.shields.io/badge/Python-3.9+-blue?style=flat-square&logo=python)](#)\n[![Jupyter](https://img.shields.io/badge/Jupyter-Notebook-F37626?style=flat-square&logo=jupyter)](#)\n\n## Objective\n> Describe the problem and goals.\n\n## Dataset\n| Feature | Type | Description |\n|---------|------|-------------|\n| col1 | int | Description |\n| col2 | str | Description |\n\n**Source:** [Link](#) | **Size:** 10K rows x 15 cols\n\n## Key Findings\n1. Finding one\n2. Finding two\n3. Finding three\n\n## Visualization\n![Chart](charts/overview.png)\n\n## Run\n```bash\npip install pandas numpy matplotlib seaborn jupyter\njupyter notebook analysis.ipynb\n```\n\n## License\nMIT\n"),
    tr("GitHub Profile"): ("<!-- Profile README -->\n<div align='center'>\n\n# Hi, I'm {name}!\n\n[![Portfolio](https://img.shields.io/badge/Portfolio-Visit-blue?style=flat-square)](#)\n[![LinkedIn](https://img.shields.io/badge/LinkedIn-Connect-0A66C2?style=flat-square&logo=linkedin)](#)\n\n</div>\n\n## About Me\n- Working on ...\n- Learning ...\n- Ask me about ...\n- your@email.com\n\n## Tech Stack\n![Python](https://img.shields.io/badge/-Python-3776AB?style=flat-square&logo=python&logoColor=white)\n![JavaScript](https://img.shields.io/badge/-JavaScript-F7DF1E?style=flat-square&logo=javascript&logoColor=black)\n\n## Stats\n<div align='center'>\n\n![Stats](https://github-readme-stats.vercel.app/api?username=USER&show_icons=true&theme=dark)\n\n</div>\n"),
    tr("Research Project"): ("# {name}\n\n[![Paper](https://img.shields.io/badge/Paper-arXiv-red?style=flat-square)](#)\n[![Venue](https://img.shields.io/badge/Venue-CVPR_2024-blue?style=flat-square)](#)\n\n**Authors:** Author 1, Author 2\n\n## Abstract\n> Write abstract here.\n\n## Contributions\n1. Contribution one\n2. Contribution two\n\n## Results\n| Method | Metric1 | Metric2 |\n|--------|---------|--------|\n| **Ours** | **95.2** | **0.12** |\n| SOTA | 93.1 | 0.18 |\n\n## Reproduce\n```bash\npip install -r requirements.txt\npython run.py --config configs/main.yaml\n```\n\n## Citation\n```bibtex\n@inproceedings{name2024,\n  title={{name}},\n  booktitle={CVPR},\n  year={2024}\n}\n```\n"),
    tr("Desktop App"): ("# {name}\n\n[![Electron](https://img.shields.io/badge/Electron-26+-47848F?style=flat-square&logo=electron)](#)\n[![Platform](https://img.shields.io/badge/Platform-Win|Mac|Linux-blue?style=flat-square)](#)\n\n## Download\n[![Windows](https://img.shields.io/badge/Windows-Download-0078D6?style=flat-square&logo=windows)](#)\n[![macOS](https://img.shields.io/badge/macOS-Download-000?style=flat-square&logo=apple)](#)\n[![Linux](https://img.shields.io/badge/Linux-Download-FCC624?style=flat-square&logo=linux)](#)\n\n## Features\n- Modern UI\n- Auto Updates\n- Portable Mode\n\n## Build\n```bash\nnpm install\nnpm run build\n```\n\n## License\nMIT\n"),
}

TAB_SNIPPETS = {
    "Screenshots": "## Screenshots\n| Preview 1 | Preview 2 |\n|---|---|\n| ![](example/1.png) | ![](example/2.png) |\n",
    "Features": "## Features\n- Performance: Lightning fast.\n- Modern UI: Clean minimal style.\n- Secure: End-to-end encryption.\n- Responsive: Works on all devices.\n",
    "Description": "## Description\nThis project streamlines your workflow and enhances productivity.\n",
    "Installation": "## Installation\n1. Clone:\n```bash\ngit clone https://github.com/user/repo.git\n```\n2. Install:\n```bash\npip install -r requirements.txt\n```\n",
    "Author": "### Author\n**Your Name** - [GitHub](https://github.com/) - [LinkedIn](#) - [Twitter](#)\n",
    "WARNING Alert": "> [!WARNING]\n> Educational purposes only. Use at your own risk!\n",
    "Note Alert": "> [!NOTE]\n> Don't forget to star the repo!\n",
    "Tip Alert": "> [!TIP]\n> Pro tip: Read the full documentation first.\n",
    "Important Alert": "> [!IMPORTANT]\n> Breaking changes in v2.0. See migration guide.\n",
    "Caution Alert": "> [!CAUTION]\n> This operation is irreversible.\n",
    "License MIT": "## License\nThis project is licensed under the **MIT License** - see [LICENSE](LICENSE) for details.\n",
    "Contributing": "## Contributing\n1. Fork the Project\n2. Create branch (`git checkout -b feature/X`)\n3. Commit (`git commit -m 'Add X'`)\n4. Push (`git push origin feature/X`)\n5. Open a Pull Request\n",
    "Usage": "## Usage\n```python\nfrom module import ClassName\n\nobj = ClassName(config='default')\nresult = obj.run()\nprint(result)\n```\n",
    "Requirements": "## Requirements\n```\nPython >= 3.8\nDependency1 >= 1.0\nDependency2 >= 2.0\n```\n",
    "GitHub Stats": "## GitHub Stats\n<div align='center'>\n\n![Stats](https://github-readme-stats.vercel.app/api?username=USER&show_icons=true&theme=dark&hide_border=true)\n\n</div>\n",
    "Roadmap": "## Roadmap\n- [x] Initial release\n- [x] Core features\n- [ ] Version 2.0\n- [ ] Mobile support\n- [ ] i18n\n",
}

# ══════════════════════════════════════════════════════════════════════════════
# EDITOR UNDO/REDO COMMAND
# ══════════════════════════════════════════════════════════════════════════════
class TextChangeCommand(QUndoCommand):
    """
    Stores a single text snapshot transition (old -> new).
    No merging — each 600 ms debounce window = one undo step.
    """
    def __init__(self, editor, old_text, new_text, cursor_pos):
        super().__init__("Text Change")
        self.editor    = editor
        self.old_text  = old_text
        self.new_text  = new_text
        self.cursor_pos = cursor_pos

    def _apply(self, text, pos):
        self.editor._undo_redo_active = True
        cur = self.editor.textCursor()
        cur.select(QTextCursor.SelectionType.Document)
        cur.insertText(text)
        cur2 = self.editor.textCursor()
        cur2.setPosition(min(pos, len(text)))
        self.editor.setTextCursor(cur2)
        self.editor._last_text = text
        self.editor._undo_redo_active = False

    def undo(self):
        self._apply(self.old_text, self.cursor_pos)

    def redo(self):
        self._apply(self.new_text, self.cursor_pos)

    def id(self):
        return -1


# ══════════════════════════════════════════════════════════════════════════════
# MARKDOWN SYNTAX HIGHLIGHTER
# ══════════════════════════════════════════════════════════════════════════════
class MarkdownHighlighter(QSyntaxHighlighter):
    def __init__(self, document):
        super().__init__(document)

    def _rules(self):
        t = T()
        def fmt(c, bold=False, italic=False):
            f = QTextCharFormat(); f.setForeground(QColor(c))
            if bold:   f.setFontWeight(QFont.Bold)
            if italic: f.setFontItalic(True)
            return f
        return [
            (re.compile(r"^# .+"),          fmt(t["syn_h1"],    bold=True)),
            (re.compile(r"^## .+"),         fmt(t["syn_h2"],    bold=True)),
            (re.compile(r"^#{3,} .+"),      fmt(t["syn_h2"])),
            (re.compile(r"\*\*[^*]+\*\*"),  fmt(t["syn_bold"],  bold=True)),
            (re.compile(r"\*[^*]+\*"),      fmt(t["syn_italic"],italic=True)),
            (re.compile(r"`[^`]+`"),        fmt(t["syn_code"])),
            (re.compile(r"^```.*"),         fmt(t["syn_fence"])),
            (re.compile(r"!\[.*?\]\(.*?\)"),fmt(t["syn_img"])),
            (re.compile(r"\[.*?\]\(.*?\)"), fmt(t["syn_link"])),
            (re.compile(r"^>.*"),           fmt(t["syn_quote"],  italic=True)),
            (re.compile(r"^\s*[-*+] "),     fmt(t["syn_list"])),
            (re.compile(r"^\s*\d+\. "),     fmt(t["syn_list"])),
            (re.compile(r"^---+$"),         fmt(t["syn_hr"])),
            (re.compile(r"<[^>]+>"),        fmt(t["syn_tag"])),
            (re.compile(r"\[[ xX]\]"),      fmt(t["syn_chk"])),
            (re.compile(r"\[!(NOTE|WARNING|TIP|IMPORTANT|CAUTION)\]"),
                                            fmt(t["syn_alert"],  bold=True)),
        ]

    def highlightBlock(self, text):
        for pattern, fmt in self._rules():
            for m in pattern.finditer(text):
                self.setFormat(m.start(), m.end()-m.start(), fmt)

# ═══════════════════════════════════════════════════════════════════════════
# LINE NUMBER AREA
# ═══════════════════════════════════════════════════════════════════════════
class LineNumberArea(QWidget):
    def __init__(self, editor):
        super().__init__(editor); self.editor = editor
    def sizeHint(self): return QSize(self.editor.line_number_area_width(), 0)
    def paintEvent(self, event): self.editor.line_number_area_paint_event(event)

# ═══════════════════════════════════════════════════════════════════════════
# FIND & REPLACE BAR
# ═══════════════════════════════════════════════════════════════════════════
class FindReplaceBar(QFrame):
    def __init__(self, editor, parent=None):
        super().__init__(parent)
        self.editor = editor
        self.setFixedHeight(50)
        lay = QHBoxLayout(self); lay.setContentsMargins(12,6,12,6); lay.setSpacing(8)
        lay.addWidget(QLabel("Find:"))
        self.find_input = QLineEdit(); self.find_input.setPlaceholderText("Find...")
        self.find_input.setFixedWidth(180); lay.addWidget(self.find_input)
        self.find_input.textChanged.connect(self.highlight_all)
        lay.addWidget(QLabel("Replace:"))
        self.replace_input = QLineEdit(); self.replace_input.setPlaceholderText("Replace with...")
        self.replace_input.setFixedWidth(180); lay.addWidget(self.replace_input)
        self.case_cb = QCheckBox("Aa"); self.case_cb.setToolTip("Case sensitive"); lay.addWidget(self.case_cb)
        self.match_lbl = QLabel(""); self.match_lbl.setTextFormat(Qt.RichText); lay.addWidget(self.match_lbl)
        for txt, fn, tip in [("Prev",self.find_prev,"Prev"),("Next",self.find_next,"Next"),
                               ("Replace",self.replace_one,"Replace"),("All",self.replace_all,"Replace all")]:
            b = QPushButton(txt); b.setToolTip(tip); b.clicked.connect(fn); lay.addWidget(b)
        lay.addStretch()
        close = QPushButton("X"); close.setFixedWidth(28)
        close.clicked.connect(self.close_bar); lay.addWidget(close)
        self.find_input.returnPressed.connect(self.find_next)
        self.apply_theme()

    def apply_theme(self):
        t = T()
        self.setStyleSheet(f"""
            QFrame      {{ background:{t['panel']}; border-top:1px solid {t['border']}; }}
            QLabel      {{ color:{t['dim']}; background:transparent; font-size:11px; }}
            QCheckBox   {{ color:{t['dim']}; background:transparent; font-size:11px; }}
            QLineEdit   {{ background:{t['input_bg']}; color:{t['editor_fg']};
                           border:1px solid {t['border']}; border-radius:6px;
                           padding:6px 10px; font-size:12px; }}
            QLineEdit:focus {{ border-color:{t['accent']}; }}
            QPushButton {{ background:{t['hover']}; color:{t['text']};
                           border:1px solid {t['border']}; border-radius:6px;
                           padding:6px 10px; font-size:11px; font-weight:600; }}
            QPushButton:hover {{ background:{t['border']}; border-color:{t['accent']}; }}
        """)

    def close_bar(self):
        self.hide(); self.editor.setExtraSelections([]); self.editor.setFocus()

    def _flags(self):
        f = QTextDocument.FindFlag(0)
        if self.case_cb.isChecked(): f |= QTextDocument.FindCaseSensitively
        return f

    def highlight_all(self):
        term = self.find_input.text()
        if not term: self.editor.setExtraSelections([]); self.match_lbl.setText(""); return
        extra = []; fmt = QTextCharFormat()
        fmt.setBackground(QColor("#3a3a00")); fmt.setForeground(QColor("#fbbf24"))
        doc = self.editor.document(); count = 0; flags = self._flags()
        cursor = doc.find(term, 0, flags)
        while not cursor.isNull():
            sel = QTextEdit.ExtraSelection(); sel.cursor = cursor; sel.format = fmt
            extra.append(sel); count += 1; cursor = doc.find(term, cursor, flags)
        self.editor.setExtraSelections(extra)
        c = "#4ade80" if count > 0 else "#f87171"
        self.match_lbl.setText(f'<span style="color:{c}">{count} found</span>')

    def find_next(self):
        term = self.find_input.text()
        if not term: return
        if not self.editor.find(term, self._flags()):
            cur = self.editor.textCursor(); cur.movePosition(QTextCursor.Start)
            self.editor.setTextCursor(cur); self.editor.find(term, self._flags())

    def find_prev(self):
        term = self.find_input.text()
        if not term: return
        flags = self._flags() | QTextDocument.FindBackward
        if not self.editor.find(term, flags):
            cur = self.editor.textCursor(); cur.movePosition(QTextCursor.End)
            self.editor.setTextCursor(cur); self.editor.find(term, flags)

    def replace_one(self):
        cur = self.editor.textCursor()
        if cur.hasSelection(): cur.insertText(self.replace_input.text())
        self.find_next(); self.highlight_all()

    def replace_all(self):
        text = self.editor.toPlainText(); term = self.find_input.text()
        repl = self.replace_input.text()
        if not term: return
        fl = 0 if self.case_cb.isChecked() else re.IGNORECASE
        new = re.sub(re.escape(term), repl, text, flags=fl)
        count = len(re.findall(re.escape(term), text, fl))
        self.editor.setPlainText(new)
        self.match_lbl.setText(f'<span style="color:#4ade80">Replaced {count}</span>')

# ═══════════════════════════════════════════════════════════════════════════
# MINIMAP
# ═══════════════════════════════════════════════════════════════════════════
class Minimap(QWidget):
    def __init__(self, editor, parent=None):
        super().__init__(parent); self.editor = editor
        self.setFixedWidth(80); self.setMouseTracking(True)
        self.setCursor(Qt.PointingHandCursor)
        self._lines = []; self._vr = 0.0; self._vp = 0.0

    def update_from_editor(self):
        self._lines = self.editor.toPlainText().split("\n")
        sb = self.editor.verticalScrollBar(); mx = max(sb.maximum(), 1)
        self._vp = sb.value() / mx
        vis = self.editor.viewport().height() / max(self.editor.fontMetrics().height(), 1)
        self._vr = min(vis / max(len(self._lines), 1), 1.0); self.update()

    def paintEvent(self, event):
        t = T(); painter = QPainter(self); w, h = self.width(), self.height()
        painter.fillRect(0, 0, w, h, QColor(t["input_bg"]))
        if not self._lines: return
        total = len(self._lines); line_h = max(h / total, 1.5)
        for i, line in enumerate(self._lines):
            y = int(i * line_h); s = line.strip()
            if not s: continue
            if   re.match(r"^#{1,2} ", s):       c = QColor(t["syn_h1"]); bh = max(line_h, 2)
            elif re.match(r"^#{3,} ", s):         c = QColor(t["syn_h2"]); bh = max(line_h, 1.5)
            elif s.startswith("```"):             c = QColor(t["syn_fence"]); bh = max(line_h, 1.5)
            elif s.startswith(">"):               c = QColor(t["syn_quote"]); bh = max(line_h, 1.5)
            elif re.match(r"^[-*+] |^\d+\. ", s):c = QColor(t["syn_list"]); bh = max(line_h, 1.5)
            elif "![" in s or "shields.io" in s: c = QColor(t["syn_img"]); bh = max(line_h, 1.5)
            else:                                 c = QColor(t["border"]); bh = max(line_h * 0.6, 1.0)
            ml = max(max((len(l) for l in self._lines if l), default=1), 1)
            bw = min(int((len(line) / ml) * (w - 8)) + 4, w - 4)
            painter.fillRect(4, y, bw, int(bh), c)
        vpy = int(self._vp * h * (1 - self._vr)); vph = max(int(self._vr * h), 20)
        ax = int(t["accent"][1:3],16); ay = int(t["accent"][3:5],16); az = int(t["accent"][5:7],16)
        painter.fillRect(0, vpy, w, vph, QColor(ax, ay, az, 50))
        painter.setPen(QPen(QColor(t["accent"]), 1)); painter.drawRect(0, vpy, w-1, vph-1)

    def mousePressEvent(self, e): self._nav(e.position().y())
    def mouseMoveEvent(self, e):
        if e.buttons() & Qt.LeftButton: self._nav(e.position().y())
    def _nav(self, y):
        sb = self.editor.verticalScrollBar()
        sb.setValue(int(y / max(self.height(), 1) * sb.maximum()))

# ═══════════════════════════════════════════════════════════════════════════
# STATS BAR
# ═══════════════════════════════════════════════════════════════════════════
class StatsBar(QFrame):
    def __init__(self, editor, parent=None):
        super().__init__(parent); self.editor = editor
        self.setFixedHeight(26); self._seps = []
        lay = QHBoxLayout(self); lay.setContentsMargins(12,0,12,0); lay.setSpacing(0)
        for attr in ["lbl_words","lbl_chars","lbl_lines","lbl_heads","lbl_code","lbl_score"]:
            lbl = QLabel(); lbl.setTextFormat(Qt.RichText); setattr(self, attr, lbl)
            lay.addWidget(lbl)
            if attr != "lbl_score":
                sep = QLabel("·"); self._seps.append(sep); lay.addWidget(sep)
        lay.addStretch()
        self.lbl_cursor = QLabel(); self.lbl_cursor.setTextFormat(Qt.RichText)
        lay.addWidget(self.lbl_cursor)
        self.editor.textChanged.connect(self.update_stats)
        self.editor.cursorPositionChanged.connect(self.update_cursor)
        self.apply_theme(); self.update_stats()

    def apply_theme(self):
        t = T()
        self.setStyleSheet(f"""
            QFrame {{ background:{t['panel']}; border-top:1px solid {t['border']}; }}
            QLabel {{ font-size:11px; font-family:'Consolas','SF Mono',monospace;
                      padding:0 4px; background:transparent; color:{t['dim']}; }}
        """)
        for sep in self._seps:
            sep.setStyleSheet(
                f"color:{t['border']};font-size:11px;background:transparent;")
        self.update_stats(); self.update_cursor()

    def _lbl(self, key, val, color=None):
        t = T(); c = color or t["dim"]
        return (f'<span style="color:{t["border"]}">{key}:</span>'
                f' <span style="color:{c}">{val}</span>')

    def update_stats(self):
        text = self.editor.toPlainText()
        words = len(text.split()) if text.strip() else 0
        chars = len(text); lines = text.count("\n") + 1 if text else 0
        headers = len(re.findall(r"^#+\s", text, re.M))
        blocks  = len(re.findall(r"^```", text, re.M)) // 2
        score, _ = calculate_readme_score(text)
        if   score >= 80: sc, sc_c = f"{score}/100", "#4ade80"
        elif score >= 50: sc, sc_c = f"{score}/100", "#fbbf24"
        else:             sc, sc_c = f"{score}/100", "#f87171"
        self.lbl_words.setText(self._lbl("Words",   f"{words:,}"))
        self.lbl_chars.setText(self._lbl("Chars",   f"{chars:,}"))
        self.lbl_lines.setText(self._lbl("Lines",   f"{lines:,}"))
        self.lbl_heads.setText(self._lbl("Headers", str(headers)))
        self.lbl_code.setText( self._lbl("Blocks",  str(blocks)))
        self.lbl_score.setText(self._lbl("Score",   sc, sc_c))

    def update_cursor(self):
        t = T(); cur = self.editor.textCursor()
        self.lbl_cursor.setText(
            f'<span style="color:{t["border"]}">Ln {cur.blockNumber()+1},'
            f' Col {cur.columnNumber()+1}</span>')

# ═══════════════════════════════════════════════════════════════════════════
# README SCORE
# ═══════════════════════════════════════════════════════════════════════════
def calculate_readme_score(text):
    score = 0; feedback = []
    if re.search(r'^#\s+', text, re.M): score += 20
    else: feedback.append("Missing main title (#)")
    if len(text) > 150: score += 15
    else: feedback.append("Add more description")
    if re.search(r'##.*install', text, re.I): score += 15
    else: feedback.append("Missing installation section")
    if re.search(r'##.*(usage|example|quick.*start)', text, re.I): score += 15
    else: feedback.append("Add usage examples")
    if re.search(r'!\[.*\]\(https://img\.shields\.io', text): score += 10
    else: feedback.append("Add badges")
    if re.search(r'```', text): score += 10
    else: feedback.append("Add code examples")
    if re.search(r'##.*license', text, re.I): score += 10
    else: feedback.append("Missing license")
    if re.search(r'##.*contribut', text, re.I): score += 5
    return score, feedback

# ═══════════════════════════════════════════════════════════════════════════
# RENDER WORKER
# ═══════════════════════════════════════════════════════════════════════════
class RenderWorker(QObject):
    finished = Signal(str)
    def __init__(self): super().__init__(); self._md=""; self._dirty=False; self._last_html=""
    def request(self, md): self._md=md; self._dirty=True
    def process(self):
        if not self._dirty: return
        self._dirty = False
        html = markdown2.markdown(self._md,
               extras=["fenced-code-blocks","tables","task_list","strike","code-friendly"])
        for old, new in [
            ("> [!NOTE]",       '<div class="markdown-alert markdown-alert-note"><b>Note</b><br>'),
            ("> [!WARNING]",    '<div class="markdown-alert markdown-alert-warning"><b>Warning</b><br>'),
            ("> [!TIP]",        '<div class="markdown-alert markdown-alert-tip"><b>Tip</b><br>'),
            ("> [!IMPORTANT]",  '<div class="markdown-alert markdown-alert-important"><b>Important</b><br>'),
            ("> [!CAUTION]",    '<div class="markdown-alert markdown-alert-caution"><b>Caution</b><br>'),
        ]: html = html.replace(old, new)
        self.finished.emit(html)

# ═══════════════════════════════════════════════════════════════════════════
# MODERN INPUT DIALOG
# ═══════════════════════════════════════════════════════════════════════════
class ModernInputDialog(QDialog):
    def __init__(self, parent=None, title="", label="", placeholder=""):
        super().__init__(parent)
        self.setFixedSize(450, 220)
        self.setWindowFlags(Qt.Popup | Qt.FramelessWindowHint)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.result_text = ""; self.ok_clicked = False
        t = T()
        outer = QFrame(self); outer.setObjectName("outer"); outer.setFixedSize(450, 220)
        lay = QVBoxLayout(outer); lay.setContentsMargins(24,20,24,20); lay.setSpacing(10)
        lay.addWidget(QLabel(title,
            styleSheet=f"color:{t['text']};font-size:16px;font-weight:600;background:transparent;"))
        lay.addWidget(QLabel(label,
            styleSheet=f"color:{t['dim']};font-size:13px;background:transparent;"))
        self.input = QLineEdit(); self.input.setPlaceholderText(placeholder)
        self.input.setStyleSheet(f"""
            QLineEdit{{background:{t['input_bg']};border:2px solid {t['border']};border-radius:8px;
                       padding:10px 14px;color:{t['editor_fg']};font-size:14px;}}
            QLineEdit:focus{{border:2px solid {t['accent']};background:{t['hover']};}}""")
        lay.addWidget(self.input)
        row = QHBoxLayout(); row.setSpacing(10)
        for txt, bg, hover, slot in [
            ("Cancel",  t['border'],  t['hover'],   self.reject),
            ("Confirm", t['accent'],  t['accent2'], self.on_confirm),
        ]:
            btn = QPushButton(txt)
            btn.setStyleSheet(f"""QPushButton{{background:{bg};color:white;border:none;
                padding:10px;border-radius:8px;font-weight:600;font-size:13px;}}
                QPushButton:hover{{background:{hover};}}""")
            btn.clicked.connect(slot); row.addWidget(btn)
        lay.addLayout(row)
        outer.setStyleSheet(
            f"QFrame#outer{{background:{t['panel']};border:1px solid {t['border']};border-radius:16px;}}")
        self.setWindowOpacity(0)

    def showEvent(self, event):
        super().showEvent(event)
        a = QPropertyAnimation(self, b"windowOpacity"); a.setDuration(200)
        a.setStartValue(0); a.setEndValue(1); a.setEasingCurve(QEasingCurve.OutCubic)
        a.start(); self._anim = a

    def on_confirm(self): self.result_text=self.input.text(); self.ok_clicked=True; self.accept()

    @staticmethod
    def get_text(parent, title, label, placeholder=""):
        dlg = ModernInputDialog(parent, title, label, placeholder)
        if dlg.exec() == QDialog.Accepted: return (dlg.result_text, dlg.ok_clicked)
        return ("", False)

# ═══════════════════════════════════════════════════════════════════════════
# EXPORT DIALOG
# ═══════════════════════════════════════════════════════════════════════════
class ExportDialog(QDialog):
    def __init__(self, markdown_text, workspace_dir, parent=None):
        super().__init__(parent)
        self.markdown_text = markdown_text; self.workspace_dir = workspace_dir
        t = T()
        self.setWindowTitle("Export README"); self.setFixedSize(460, 300)
        self.setStyleSheet(f"""
            QDialog     {{ background:{t['panel']}; color:{t['text']}; }}
            QLabel      {{ color:{t['text']}; font-size:13px; background:transparent; }}
            QPushButton {{ background:{t['hover']}; color:{t['editor_fg']};
                           border:2px solid {t['border']}; border-radius:10px;
                           padding:16px 10px; font-weight:600; font-size:12px; }}
            QPushButton:hover  {{ background:{t['border']}; border-color:{t['accent']}; color:white; }}
            QPushButton:pressed{{ background:{t['accent']}; border-color:{t['accent']}; }}
        """)
        lay = QVBoxLayout(self); lay.setContentsMargins(24,20,24,20); lay.setSpacing(14)
        lay.addWidget(QLabel("Export As",
            styleSheet=f"font-size:18px;font-weight:700;color:{t['text']};background:transparent;"))
        lay.addWidget(QLabel("Choose your export format:",
            styleSheet=f"color:{t['dim']};font-size:12px;background:transparent;"))
        grid = QGridLayout(); grid.setSpacing(10)
        for i, (label, tip, fn) in enumerate([
            ("HTML",    "Styled web page",         self.export_html),
            ("Markdown","Clean .md with metadata", self.export_md),
            ("PDF",     "Browser -> Print -> PDF",   self.export_pdf),
            ("DOCX",    "Word document",            self.export_docx),
        ]):
            btn = QPushButton(label); btn.setToolTip(tip)
            btn.setMinimumHeight(60); btn.clicked.connect(fn)
            grid.addWidget(btn, i//2, i%2)
        lay.addLayout(grid)
        lay.addWidget(QLabel("PDF: browser -> Ctrl+P -> Save as PDF",
            styleSheet=f"color:{t['dim']};font-size:10px;background:transparent;"))

    def _html(self):
        t = T(); body = markdown2.markdown(self.markdown_text,
            extras=["fenced-code-blocks","tables","task_list","strike"])
        css = _preview_css(t)
        return f"<!DOCTYPE html><html lang='en'><head><meta charset='UTF-8'><title>README</title>{css}</head><body>{body}</body></html>"

    def export_html(self):
        p, _ = QFileDialog.getSaveFileName(self, "Export HTML", "README.html", "HTML (*.html)")
        if not p: return
        try: Path(p).write_text(self._html(), encoding="utf-8"); QMessageBox.information(self,"Done",f"Saved:\n{p}"); self.accept()
        except Exception as e: QMessageBox.critical(self,"Error",str(e))

    def export_md(self):
        p, _ = QFileDialog.getSaveFileName(self, "Export Markdown", "README.md", "Markdown (*.md)")
        if not p: return
        try:
            Path(p).write_text(f"<!-- README Builder Pro | {datetime.now():%Y-%m-%d} -->\n\n{self.markdown_text}", encoding="utf-8")
            QMessageBox.information(self,"Done",f"Saved:\n{p}"); self.accept()
        except Exception as e: QMessageBox.critical(self,"Error",str(e))

    def export_pdf(self):
        import tempfile, webbrowser
        try:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".html", mode="w", encoding="utf-8")
            tmp.write(self._html()); tmp.close()
            webbrowser.open(f"file://{tmp.name}")
            QMessageBox.information(self,"PDF","Opened in browser.\n1. Ctrl+P  2. Save as PDF  3. Save")
            self.accept()
        except Exception as e: QMessageBox.critical(self,"Error",str(e))

    def export_docx(self):
        p, _ = QFileDialog.getSaveFileName(self, "Export DOCX", "README.docx", "Word (*.docx)")
        if not p: return
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            doc = Document(); doc.styles["Normal"].font.name="Calibri"
            doc.styles["Normal"].font.size = Pt(11)
            lines = self.markdown_text.split("\n"); in_code = False; code_buf = []
            for line in lines:
                if line.startswith("```"):
                    if in_code:
                        p2 = doc.add_paragraph(); r = p2.add_run("\n".join(code_buf))
                        r.font.name="Courier New"; r.font.size=Pt(9)
                        r.font.color.rgb=RGBColor(0x48,0x5A,0x6A); code_buf=[]; in_code=False
                    else: in_code=True
                    continue
                if in_code: code_buf.append(line); continue
                if   line.startswith("# "):  doc.add_heading(line[2:],  1)
                elif line.startswith("## "): doc.add_heading(line[3:],  2)
                elif line.startswith("### "):doc.add_heading(line[4:],  3)
                elif line.startswith(("- ","* ")): doc.add_paragraph(line[2:],style="List Bullet")
                elif re.match(r"^\d+\. ",line): doc.add_paragraph(re.sub(r"^\d+\. ","",line),style="List Number")
                elif line.strip() in ("","---"): doc.add_paragraph()
                else:
                    clean = re.sub(r"\*\*([^*]+)\*\*",r"\1",line)
                    clean = re.sub(r"\*([^*]+)\*",r"\1",clean)
                    clean = re.sub(r"`([^`]+)`",r"\1",clean)
                    clean = re.sub(r"\[([^\]]+)\]\([^\)]+\)",r"\1",clean)
                    if clean.strip(): doc.add_paragraph(clean)
            doc.save(p)
            QMessageBox.information(self,"Done",f"Saved:\n{p}"); self.accept()
        except ImportError: QMessageBox.warning(self,"Missing","pip install python-docx")
        except Exception as e: QMessageBox.critical(self,"Error",str(e))


# ═══════════════════════════════════════════════════════════════════════════
# DRAWING TOOL ITEMS
# ═══════════════════════════════════════════════════════════════════════════
class EditableTextItem(QGraphicsTextItem):
    def __init__(self, text="Text", parent=None):
        super().__init__(parent)
        self.setPlainText(text)
        self.setFlags(QGraphicsItem.ItemIsMovable | QGraphicsItem.ItemIsSelectable | QGraphicsItem.ItemIsFocusable)
        self.setTextInteractionFlags(Qt.NoTextInteraction)
        self.setDefaultTextColor(QColor("#fff")); self.setFont(QFont("Arial", 14))
    def mouseDoubleClickEvent(self, event):
        self.setTextInteractionFlags(Qt.TextEditorInteraction); self.setFocus(); super().mouseDoubleClickEvent(event)
    def focusOutEvent(self, event):
        self.setTextInteractionFlags(Qt.NoTextInteraction); super().focusOutEvent(event)

class ResizeHandle(QGraphicsRectItem):
    def __init__(self, parent, pos_index):
        super().__init__(parent)
        self.parent_item = parent; self.pos_index = pos_index
        self._in_update = False
        self.setRect(-6, -6, 12, 12); self.setBrush(QBrush(QColor("#8b5cf6")))
        self.setPen(QPen(QColor("#ffffff"), 1)); self.setZValue(1000)
        self.setFlags(QGraphicsItem.ItemIsMovable | QGraphicsItem.ItemIsSelectable |
                      QGraphicsItem.ItemSendsGeometryChanges | QGraphicsItem.ItemIgnoresTransformations)
        cursors = [Qt.SizeFDiagCursor, Qt.SizeBDiagCursor, Qt.SizeBDiagCursor, Qt.SizeFDiagCursor]
        self.setCursor(cursors[pos_index]); self.update_position()

    def update_position(self):
        if self._in_update: return
        self._in_update = True
        r = self.parent_item.rect()
        if   self.pos_index == 0: self.setPos(r.topLeft())
        elif self.pos_index == 1: self.setPos(r.topRight())
        elif self.pos_index == 2: self.setPos(r.bottomLeft())
        elif self.pos_index == 3: self.setPos(r.bottomRight())
        self._in_update = False

    def itemChange(self, change, value):
        if change == QGraphicsItem.ItemPositionChange and self.parent_item.isSelected() and not self._in_update:
            self._in_update = True
            self.parent_item.handle_move(self.pos_index, value)
            self._in_update = False
        return super().itemChange(change, value)

class ResizableRectItem(QGraphicsRectItem):
    def __init__(self, rect, pen, brush):
        super().__init__()
        self.setRect(rect); self.setPen(pen); self.setBrush(brush)
        self.setFlags(QGraphicsItem.ItemIsMovable | QGraphicsItem.ItemIsSelectable)
        self.handles = [ResizeHandle(self, i) for i in range(4)]
        for h in self.handles: h.hide()

    def paint(self, painter, option, widget):
        super().paint(painter, option, widget)
        if self.isSelected():
            for h in self.handles: h.show()
        else:
            for h in self.handles: h.hide()

    def handle_move(self, index, pos):
        r = self.rect()
        if index == 0: r.setTopLeft(pos)
        elif index == 1: r.setTopRight(pos)
        elif index == 2: r.setBottomLeft(pos)
        elif index == 3: r.setBottomRight(pos)
        self.setRect(r.normalized())
        for h in self.handles: h.update_position()

class ResizableEllipseItem(QGraphicsEllipseItem):
    def __init__(self, rect, pen, brush):
        super().__init__()
        self.setRect(rect); self.setPen(pen); self.setBrush(brush)
        self.setFlags(QGraphicsItem.ItemIsMovable | QGraphicsItem.ItemIsSelectable)
        self.handles = [ResizeHandle(self, i) for i in range(4)]
        for h in self.handles: h.hide()

    def paint(self, painter, option, widget):
        super().paint(painter, option, widget)
        if self.isSelected():
            for h in self.handles: h.show()
        else:
            for h in self.handles: h.hide()

    def handle_move(self, index, pos):
        r = self.rect()
        if index == 0: r.setTopLeft(pos)
        elif index == 1: r.setTopRight(pos)
        elif index == 2: r.setBottomLeft(pos)
        elif index == 3: r.setBottomRight(pos)
        self.setRect(r.normalized())
        for h in self.handles: h.update_position()

class ResizablePixmapItem(QGraphicsPixmapItem):
    def __init__(self, pixmap):
        super().__init__(pixmap)
        self.setFlags(QGraphicsItem.ItemIsMovable | QGraphicsItem.ItemIsSelectable)
        self._rect = QRectF(0, 0, pixmap.width(), pixmap.height())
        self.handles = [ResizeHandle(self, i) for i in range(4)]
        for h in self.handles: h.hide()

    def rect(self): return self._rect
    def setRect(self, r):
        self.prepareGeometryChange(); self._rect = r
        for h in self.handles: h.update_position()

    def boundingRect(self): return self._rect

    def paint(self, painter, option, widget):
        painter.drawPixmap(self._rect.toRect(), self.pixmap())
        if self.isSelected():
            for h in self.handles: h.show()
        else:
            for h in self.handles: h.hide()

    def handle_move(self, index, pos):
        r = self.rect()
        if index == 0: r.setTopLeft(pos)
        elif index == 1: r.setTopRight(pos)
        elif index == 2: r.setBottomLeft(pos)
        elif index == 3: r.setBottomRight(pos)
        self.setRect(r.normalized())

# ═══════════════════════════════════════════════════════════════════════════
# DRAWING TOOL  — fixed zoom + Alt+drag copy
# ═══════════════════════════════════════════════════════════════════════════
class DrawingTool(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Drawing Tool"); self.setMinimumSize(1000, 700)
        self.setStyleSheet("background:#18181b;")
        self.current_tool = "select"; self.current_color = QColor("#8b5cf6")
        self.fill_color = QColor(Qt.transparent); self.stroke_width = 2; self.text_size = 14
        self.drawing = False; self.start_pos = None; self.current_item = None
        self.zoom_level = 1.0; self.undo_stack = []; self.redo_stack = []
        self.canvas_color = QColor("#181818")
        self._alt_copy_source = None
        self._grid_items = []       # grid line items
        self._grid_visible = True   # grid toggle state
        self._right_press_pos = None  # for right-click hold detection
        self._right_press_item = None
        layout = QVBoxLayout(self); layout.setContentsMargins(0,0,0,0); layout.setSpacing(0)
        layout.addWidget(self.create_toolbar())
        self.scene = QGraphicsScene(); self.scene.setSceneRect(0,0,1400,900)
        self.scene.setBackgroundBrush(QBrush(QColor("#181818")))
        self.view = QGraphicsView(self.scene)
        self.view.setRenderHint(QPainter.Antialiasing); self.view.setRenderHint(QPainter.SmoothPixmapTransform)
        self.view.setStyleSheet("border:none;"); self.view.setDragMode(QGraphicsView.NoDrag)
        self.view.setTransformationAnchor(QGraphicsView.AnchorUnderMouse)
        self.view.setResizeAnchor(QGraphicsView.AnchorUnderMouse)
        self._draw_grid()
        layout.addWidget(self.view)
        self.view.viewport().installEventFilter(self)
        # Zoom label
        self.zoom_lbl = QLabel("100%", self)
        self.zoom_lbl.setStyleSheet("color:#9ca3af;font-size:11px;background:transparent;")
        self.zoom_lbl.move(10, self.height()-30)

    def create_toolbar(self):
        toolbar = QFrame(); toolbar.setFixedHeight(130)
        toolbar.setStyleSheet("""
            QFrame{background:#1e1e24;border-bottom:1px solid #3f3f46;}
            QToolButton{background:transparent;color:#9ca3af;border:2px solid transparent;border-radius:6px;padding:5px 7px;font-weight:600;font-size:11px;}
            QToolButton:hover{background:#2a2a32;color:#fff;}
            QToolButton:checked{background:#8b5cf6;color:#fff;border-color:#8b5cf6;}
            QPushButton{background:#27272a;color:#fff;border:1px solid #3f3f46;border-radius:6px;padding:5px 10px;font-weight:600;font-size:11px;}
            QPushButton:hover{background:#3f3f46;}
            QSpinBox{background:#27272a;color:#fff;border:1px solid #3f3f46;border-radius:6px;padding:4px;font-size:11px;width:40px;}
            QLabel{color:#9ca3af;font-size:10px;background:transparent;}
        """)
        main_lay = QVBoxLayout(toolbar); main_lay.setContentsMargins(10,6,10,6); main_lay.setSpacing(5)
        row1 = QHBoxLayout(); row2 = QHBoxLayout(); row3 = QHBoxLayout()
        main_lay.addLayout(row1); main_lay.addLayout(row2); main_lay.addLayout(row3)

        # Row 1: Tools with keyboard shortcuts shown
        self.tool_buttons = {}
        for name, label, key in [
            ("select", "Select [V]", "V"),
            ("rect",   "Rect [R]",   "R"),
            ("ellipse","Circle [E]", "E"),
            ("line",   "Line [L]",   "L"),
            ("arrow",  "Arrow [A]",  "A"),
            ("pen",    "Pen [P]",    "P"),
            ("text",   "Text [T]",   "T"),
        ]:
            btn = QToolButton(); btn.setText(label); btn.setCheckable(True)
            btn.setToolTip(f"Tool: {name}  |  Shortcut: {key}")
            btn.clicked.connect(lambda _, n=name: self.set_tool(n))
            row1.addWidget(btn); self.tool_buttons[name] = btn
        self.tool_buttons["select"].setChecked(True)
        row1.addWidget(self._sep())
        for label, slot, tip in [
            ("Image [I]", self.import_image, "Import image  [I]"),
            ("Table",     self.insert_table, "Insert markdown table"),
            ("Font [F]",  self.choose_font,  "Choose font  [F]"),
        ]:
            b = QPushButton(label); b.setToolTip(tip); b.clicked.connect(slot); row1.addWidget(b)
        row1.addWidget(QLabel("Size:"))
        self.text_spin = QSpinBox(); self.text_spin.setRange(8,72); self.text_spin.setValue(14)
        self.text_spin.valueChanged.connect(self.change_text_size); row1.addWidget(self.text_spin)
        self.zoom_display = QLabel("Zoom: 100%"); row1.addWidget(self.zoom_display)
        row1.addStretch()

        # Row 2: Colors + Stroke + Undo/Redo + Zoom
        for label, slot, tip in [
            ("Stroke Color", self.choose_color,       "Set stroke color  [C]"),
            ("Fill Color",   self.choose_fill_color,  "Set fill color  [Shift+C]"),
        ]:
            b = QPushButton(label); b.setToolTip(tip); b.clicked.connect(slot); row2.addWidget(b)
        row2.addWidget(QLabel("Stroke:"))
        self.stroke_spin = QSpinBox(); self.stroke_spin.setRange(1,20); self.stroke_spin.setValue(2)
        self.stroke_spin.valueChanged.connect(lambda v: setattr(self,'stroke_width',v)); row2.addWidget(self.stroke_spin)
        row2.addWidget(self._sep())
        for label, slot, tip in [
            ("Undo [Ctrl+Z]", self.undo, "Undo  Ctrl+Z"),
            ("Redo [Ctrl+Y]", self.redo, "Redo  Ctrl+Y"),
        ]:
            b = QPushButton(label); b.setToolTip(tip); b.clicked.connect(slot); row2.addWidget(b)
        row2.addWidget(self._sep())
        row2.addWidget(QLabel("Zoom:"))
        for label, tip, fn in [
            ("- [Ctrl+-]", "Zoom out  Ctrl+-", lambda: self.zoom(-0.15)),
            ("+ [Ctrl++]", "Zoom in   Ctrl++", lambda: self.zoom(0.15)),
            ("Reset [0]",  "Reset zoom  [0]",  self.reset_zoom),
        ]:
            b = QPushButton(label); b.setToolTip(tip); b.clicked.connect(fn); row2.addWidget(b)
        row2.addStretch()

        # Row 3: Backgrounds + Delete + Export
        row3.addWidget(QLabel("Background:"))
        for label, color, tip in [
            ("Dark",    "#181818", "Dark background"),
            ("Black",   "#000000", "Pure black"),
            ("White",   "#ffffff", "Pure white"),
            ("Paper",   "#f5f0e8", "Paper / cream"),
            ("Grid",    None,      "Grid overlay on dark"),
            ("Custom",  "custom",  "Choose custom color"),
        ]:
            b = QPushButton(label); b.setToolTip(tip)
            b.clicked.connect(lambda _, c=color: self._set_bg(c))
            row3.addWidget(b)
        # Grid toggle button
        self.grid_toggle_btn = QPushButton("Grid: ON")
        self.grid_toggle_btn.setToolTip("Toggle grid visibility")
        self.grid_toggle_btn.setStyleSheet("background:#27272a;color:#4ade80;border:1px solid #3f3f46;border-radius:6px;padding:5px 10px;font-weight:600;font-size:11px;")
        self.grid_toggle_btn.clicked.connect(self.toggle_grid)
        row3.addWidget(self.grid_toggle_btn)
        row3.addWidget(self._sep())
        for label, slot, tip, style in [
            ("Del [Delete]",   self.delete_selected, "Delete selected  [Delete]", "background:#ef4444;"),
            ("Export PNG [S]", self.export_png,       "Export to PNG  [S]",       "background:#8b5cf6;"),
        ]:
            b = QPushButton(label); b.setToolTip(tip); b.setStyleSheet(style)
            b.clicked.connect(slot); row3.addWidget(b)
        row3.addStretch()
        return toolbar

    def _set_bg(self, color):
        """Set canvas background. color=None draws grid, color='custom' opens dialog."""
        if color == "custom":
            c = QColorDialog.getColor(self.canvas_color, self)
            if c.isValid(): color = c.name()
            else: return
        if color is None:
            # Grid mode — dark bg + draw grid
            self.canvas_color = QColor("#181818")
            self.scene.setBackgroundBrush(QBrush(self.canvas_color))
            self._grid_visible = True
            self._draw_grid()
            if hasattr(self, 'grid_toggle_btn'):
                self.grid_toggle_btn.setText("Grid: ON")
                self.grid_toggle_btn.setStyleSheet("background:#27272a;color:#4ade80;border:1px solid #3f3f46;border-radius:6px;padding:5px 10px;font-weight:600;font-size:11px;")
        else:
            self.canvas_color = QColor(color)
            self.scene.setBackgroundBrush(QBrush(self.canvas_color))

    def _draw_grid(self):
        # Remove old grid lines first
        for item in self._grid_items:
            self.scene.removeItem(item)
        self._grid_items.clear()
        pen = QPen(QColor("#21262d"), 1, Qt.DotLine)
        for x in range(0, 1400, 50):
            line = self.scene.addLine(x, 0, x, 900, pen)
            line.setZValue(-1000)   # always behind everything
            self._grid_items.append(line)
        for y in range(0, 900, 50):
            line = self.scene.addLine(0, y, 1400, y, pen)
            line.setZValue(-1000)
            self._grid_items.append(line)

    def toggle_grid(self):
        self._grid_visible = not self._grid_visible
        for item in self._grid_items:
            item.setVisible(self._grid_visible)
        if hasattr(self, 'grid_toggle_btn'):
            if self._grid_visible:
                self.grid_toggle_btn.setText("Grid: ON")
                self.grid_toggle_btn.setStyleSheet("background:#27272a;color:#4ade80;border:1px solid #3f3f46;border-radius:6px;padding:5px 10px;font-weight:600;font-size:11px;")
            else:
                self.grid_toggle_btn.setText("Grid: OFF")
                self.grid_toggle_btn.setStyleSheet("background:#27272a;color:#6b7280;border:1px solid #3f3f46;border-radius:6px;padding:5px 10px;font-weight:600;font-size:11px;")

    def _sep(self):
        sep = QFrame(); sep.setFrameShape(QFrame.VLine)
        sep.setStyleSheet("background:#3f3f46;"); sep.setFixedSize(1,40); return sep

    def set_tool(self, tool):
        self.current_tool = tool
        for name, btn in self.tool_buttons.items(): btn.setChecked(name == tool)
        self.view.setDragMode(QGraphicsView.RubberBandDrag if tool == "select" else QGraphicsView.NoDrag)

    def choose_canvas_color(self):
        self._set_bg("custom")

    def choose_color(self):
        color = QColorDialog.getColor(self.current_color, self)
        if color.isValid():
            self.current_color = color
            for item in self.scene.selectedItems():
                if hasattr(item,'setPen'): p=item.pen(); p.setColor(color); item.setPen(p)
                elif hasattr(item,'setDefaultTextColor'): item.setDefaultTextColor(color)

    def choose_fill_color(self):
        color = QColorDialog.getColor(self.fill_color, self)
        if color.isValid():
            self.fill_color = color
            for item in self.scene.selectedItems():
                if hasattr(item,'setBrush'): item.setBrush(QBrush(color))

    def change_text_size(self, val):
        self.text_size = val
        for item in self.scene.selectedItems():
            if isinstance(item, EditableTextItem): f=item.font(); f.setPointSize(val); item.setFont(f)

    def import_image(self):
        path, _ = QFileDialog.getOpenFileName(self,"Import Image","","Images (*.png *.jpg *.jpeg *.bmp *.gif)")
        if path:
            item = ResizablePixmapItem(QPixmap(path))
            self.scene.addItem(item)
            item.setPos(self.view.mapToScene(self.view.viewport().rect().center()))
            self._save_action(item,'add')

    def insert_table(self):
        item = EditableTextItem(" \n | Header 1 | Header 2 | \n |---|---| \n | Cell 1 | Cell 2 | \n ")
        item.setDefaultTextColor(self.current_color)
        item.setPos(self.view.mapToScene(self.view.viewport().rect().center()))
        f = item.font(); f.setPointSize(self.text_size); item.setFont(f)
        self.scene.addItem(item); self._save_action(item,'add')

    def choose_font(self):
        font, ok = QFontDialog.getFont(QFont("Arial",14), self)
        if ok:
            for item in self.scene.selectedItems():
                if isinstance(item, EditableTextItem): item.setFont(font)

    def zoom(self, delta):
        self.zoom_level = max(0.05, min(self.zoom_level + delta, 5.0))
        self.view.resetTransform(); self.view.scale(self.zoom_level, self.zoom_level)
        pct = int(self.zoom_level * 100)
        self.zoom_display.setText(f"Zoom: {pct}%")

    def reset_zoom(self):
        self.zoom_level = 1.0; self.view.resetTransform()
        self.zoom_display.setText("Zoom: 100%")

    def delete_selected(self):
        for item in self.scene.selectedItems():
            self._save_action(item,'delete'); item.hide(); item.setSelected(False)

    def keyPressEvent(self, e):
        k = e.key(); mod = e.modifiers()
        if   k == Qt.Key_Z and mod & Qt.ControlModifier:  self.undo()
        elif k == Qt.Key_Y and mod & Qt.ControlModifier:  self.redo()
        elif k == Qt.Key_Plus  and mod & Qt.ControlModifier: self.zoom(0.15)
        elif k == Qt.Key_Equal and mod & Qt.ControlModifier: self.zoom(0.15)
        elif k == Qt.Key_Minus and mod & Qt.ControlModifier: self.zoom(-0.15)
        elif k == Qt.Key_0:   self.reset_zoom()
        elif k in (Qt.Key_Delete, Qt.Key_Backspace): self.delete_selected()
        elif k == Qt.Key_V and not (mod & Qt.ControlModifier): self.set_tool("select")
        elif k == Qt.Key_R:   self.set_tool("rect")
        elif k == Qt.Key_E:   self.set_tool("ellipse")
        elif k == Qt.Key_L:   self.set_tool("line")
        elif k == Qt.Key_A and not (mod & Qt.ControlModifier): self.set_tool("arrow")
        elif k == Qt.Key_P and not (mod & Qt.ControlModifier): self.set_tool("pen")
        elif k == Qt.Key_T and not (mod & Qt.ControlModifier): self.set_tool("text")
        elif k == Qt.Key_F and not (mod & Qt.ControlModifier): self.choose_font()
        elif k == Qt.Key_I and not (mod & Qt.ControlModifier): self.import_image()
        elif k == Qt.Key_C and not (mod & Qt.ControlModifier):
            if mod & Qt.ShiftModifier: self.choose_fill_color()
            else: self.choose_color()
        elif k == Qt.Key_S and not (mod & Qt.ControlModifier): self.export_png()
        else: super().keyPressEvent(e)

    def _save_action(self, item, action_type):
        self.undo_stack.append({'item':item,'type':action_type}); self.redo_stack.clear()

    def undo(self):
        if self.undo_stack:
            a = self.undo_stack.pop()
            a['item'].hide() if a['type']=='add' else a['item'].show()
            self.redo_stack.append(a)

    def redo(self):
        if self.redo_stack:
            a = self.redo_stack.pop()
            a['item'].show() if a['type']=='add' else a['item'].hide()
            self.undo_stack.append(a)

    def _clone_item(self, item):
        """Clone a graphics item for Alt+drag copy."""
        if isinstance(item, ResizableRectItem):
            clone = ResizableRectItem(item.rect(), item.pen(), item.brush())
            clone.setPos(item.pos())
            return clone
        elif isinstance(item, ResizableEllipseItem):
            clone = ResizableEllipseItem(item.rect(), item.pen(), item.brush())
            clone.setPos(item.pos())
            return clone
        elif isinstance(item, EditableTextItem):
            clone = EditableTextItem(item.toPlainText())
            clone.setDefaultTextColor(item.defaultTextColor())
            clone.setFont(item.font())
            clone.setPos(item.pos())
            return clone
        elif isinstance(item, ResizablePixmapItem):
            clone = ResizablePixmapItem(item.pixmap())
            clone.setPos(item.pos())
            return clone
        elif isinstance(item, QGraphicsPathItem):
            clone = QGraphicsPathItem(item.path())
            clone.setPen(item.pen())
            clone.setPos(item.pos())
            clone.setFlags(QGraphicsItem.ItemIsMovable | QGraphicsItem.ItemIsSelectable)
            return clone
        elif isinstance(item, QGraphicsLineItem):
            clone = self.scene.addLine(item.line(), item.pen())
            clone.setPos(item.pos())
            clone.setFlags(QGraphicsItem.ItemIsMovable | QGraphicsItem.ItemIsSelectable)
            return clone
        return None

    def _show_item_context_menu(self, item, global_pos):
        """Right-click context menu: scale, duplicate, delete, color."""
        ms = """
            QMenu { background:#1e1e24; border:1px solid #3f3f46; border-radius:10px;
                    padding:6px 4px; color:#e4e4e7; font-size:12px; }
            QMenu::item { padding:8px 24px 8px 16px; border-radius:6px; margin:2px 4px; }
            QMenu::item:selected { background:#8b5cf6; color:#fff; }
            QMenu::separator { height:1px; background:#3f3f46; margin:4px 10px; }
        """
        menu = QMenu(self); menu.setStyleSheet(ms)

        # Scale submenu
        scale_menu = QMenu("Scale", menu); scale_menu.setStyleSheet(ms)
        for pct, label in [(50,"50%"), (75,"75%"), (100,"100% (reset)"), (125,"125%"),
                           (150,"150%"), (200,"200%"), (300,"300%")]:
            scale_menu.addAction(label).triggered.connect(
                lambda _, p=pct, it=item: self._scale_item(it, p / 100.0))
        menu.addMenu(scale_menu)

        # Flip submenu
        flip_menu = QMenu("Flip", menu); flip_menu.setStyleSheet(ms)
        flip_menu.addAction("Flip Horizontal").triggered.connect(
            lambda _, it=item: it.setTransform(QTransform().scale(-1, 1), True))
        flip_menu.addAction("Flip Vertical").triggered.connect(
            lambda _, it=item: it.setTransform(QTransform().scale(1, -1), True))
        menu.addMenu(flip_menu)

        menu.addSeparator()

        # Duplicate
        menu.addAction("Duplicate  Alt+Drag").triggered.connect(
            lambda _, it=item: self._duplicate_item(it))

        # Bring to front / Send to back
        menu.addAction("Bring to Front").triggered.connect(
            lambda _, it=item: it.setZValue(it.zValue() + 1))
        menu.addAction("Send to Back").triggered.connect(
            lambda _, it=item: it.setZValue(it.zValue() - 1))

        menu.addSeparator()

        # Delete
        act_del = menu.addAction("Delete")
        act_del.setEnabled(True)
        act_del.triggered.connect(lambda _, it=item: (
            self._save_action(it, 'delete'), it.hide(), it.setSelected(False)))

        menu.exec(global_pos)

    def _scale_item(self, item, factor):
        """Scale item around its center."""
        if hasattr(item, 'rect'):
            r = item.rect()
            cx, cy = r.center().x(), r.center().y()
            new_w  = max(r.width()  * factor, 5)
            new_h  = max(r.height() * factor, 5)
            from PySide6.QtCore import QRectF
            new_r = QRectF(cx - new_w/2, cy - new_h/2, new_w, new_h)
            item.setRect(new_r)
            if hasattr(item, 'handles'):
                for h in item.handles: h.update_position()
        else:
            # For line/path/text items use QTransform scale
            t = item.transform()
            item.setTransform(QTransform().scale(factor, factor), True)

    def _duplicate_item(self, item):
        clone = self._clone_item(item)
        if clone:
            self.scene.addItem(clone)
            clone.setPos(item.pos() + QPointF(20, 20))
            item.setSelected(False)
            clone.setSelected(True)
            self._save_action(clone, 'add')

    def eventFilter(self, obj, event):
        if obj == self.view.viewport():
            # Wheel zoom
            if event.type() == event.Type.Wheel:
                delta = event.angleDelta().y()
                factor = 1.12 if delta > 0 else 0.89
                self.zoom_level = max(0.05, min(self.zoom_level * factor, 5.0))
                self.view.resetTransform(); self.view.scale(self.zoom_level, self.zoom_level)
                self.zoom_display.setText(f"Zoom: {int(self.zoom_level*100)}%")
                return True

            # Right-click context menu on item
            if (event.type() == event.Type.MouseButtonPress
                    and event.button() == Qt.RightButton):
                pos = self.view.mapToScene(event.position().toPoint())
                item_at = self.scene.itemAt(pos, self.view.transform())
                # Filter out grid lines and handles
                if item_at and not isinstance(item_at, ResizeHandle) and item_at not in self._grid_items:
                    self._show_item_context_menu(item_at, event.globalPosition().toPoint())
                    return True

            # Alt+drag copy
            if (event.type() == event.Type.MouseButtonPress
                    and event.modifiers() & Qt.AltModifier
                    and self.current_tool == "select"):
                pos = self.view.mapToScene(event.position().toPoint())
                item_at = self.scene.itemAt(pos, self.view.transform())
                if item_at and not isinstance(item_at, ResizeHandle):
                    clone = self._clone_item(item_at)
                    if clone:
                        self.scene.addItem(clone)
                        clone.setPos(item_at.pos() + QPointF(20, 20))
                        # Deselect original, select clone
                        item_at.setSelected(False)
                        clone.setSelected(True)
                        self._save_action(clone, 'add')
                        return True

            if event.type() == event.Type.MouseButtonPress and self.current_tool != "select":
                self.start_pos = self.view.mapToScene(event.position().toPoint())
                item_at = self.scene.itemAt(self.start_pos, self.view.transform())
                if item_at and not isinstance(item_at, QGraphicsLineItem): return super().eventFilter(obj, event)
                self.drawing = True
                if self.current_tool == "text":
                    item = EditableTextItem("Double-click to edit")
                    item.setDefaultTextColor(self.current_color); item.setPos(self.start_pos)
                    f = item.font(); f.setPointSize(self.text_size); item.setFont(f)
                    self.scene.addItem(item); self.drawing = False; self._save_action(item,'add'); return True
                elif self.current_tool == "pen":
                    path = QPainterPath(); path.moveTo(self.start_pos)
                    pen = QPen(self.current_color, self.stroke_width, Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin)
                    self.current_item = self.scene.addPath(path, pen)
            elif event.type() == event.Type.MouseMove and self.drawing and self.current_tool != "select":
                if not self.start_pos: return False
                pos = self.view.mapToScene(event.position().toPoint())
                pen = QPen(self.current_color, self.stroke_width); brush = QBrush(self.fill_color)
                if self.current_tool == "rect":
                    if self.current_item and self.current_item in self.scene.items(): self.scene.removeItem(self.current_item)
                    self.current_item = ResizableRectItem(QRectF(self.start_pos, pos).normalized(), pen, brush)
                    self.scene.addItem(self.current_item)
                elif self.current_tool == "ellipse":
                    if self.current_item and self.current_item in self.scene.items(): self.scene.removeItem(self.current_item)
                    self.current_item = ResizableEllipseItem(QRectF(self.start_pos, pos).normalized(), pen, brush)
                    self.scene.addItem(self.current_item)
                elif self.current_tool == "line":
                    if self.current_item and self.current_item in self.scene.items(): self.scene.removeItem(self.current_item)
                    self.current_item = self.scene.addLine(QLineF(self.start_pos, pos), pen)
                    self.current_item.setFlags(QGraphicsItem.ItemIsMovable | QGraphicsItem.ItemIsSelectable)
                elif self.current_tool == "arrow":
                    import math
                    if self.current_item and self.current_item in self.scene.items(): self.scene.removeItem(self.current_item)
                    path = QPainterPath(); path.moveTo(self.start_pos); path.lineTo(pos)
                    line = QLineF(self.start_pos, pos); angle = line.angle(); arrow_size = 15
                    p1 = pos - QPointF(math.cos(math.radians(angle+30))*arrow_size, -math.sin(math.radians(angle+30))*arrow_size)
                    p2 = pos - QPointF(math.cos(math.radians(angle-30))*arrow_size, -math.sin(math.radians(angle-30))*arrow_size)
                    path.moveTo(p1); path.lineTo(pos); path.lineTo(p2)
                    self.current_item = self.scene.addPath(path, pen)
                    self.current_item.setFlags(QGraphicsItem.ItemIsMovable | QGraphicsItem.ItemIsSelectable)
                elif self.current_tool == "pen":
                    if self.current_item:
                        path = self.current_item.path(); path.lineTo(pos); self.current_item.setPath(path)
            elif event.type() == event.Type.MouseButtonRelease and self.current_tool != "select":
                if self.drawing and self.current_item: self._save_action(self.current_item,'add')
                self.drawing = False; self.current_item = None
        return super().eventFilter(obj, event)

    def export_png(self):
        path, _ = QFileDialog.getSaveFileName(self,"Save Drawing","","PNG (*.png)")
        if path:
            rect = self.scene.itemsBoundingRect()
            img = QImage(rect.size().toSize(), QImage.Format_ARGB32); img.fill(QColor("#181818"))
            painter = QPainter(img); painter.setRenderHint(QPainter.Antialiasing)
            self.scene.render(painter, QRectF(), rect); painter.end(); img.save(path)
            QMessageBox.information(self,"Success",f"Saved to {path}")


# ══════════════════════════════════════════════════════════════════════════════
# SMART DIALOGS
# ══════════════════════════════════════════════════════════════════════════════
class ImagesDialog(QDialog):
    insert_code = Signal(str)
    def __init__(self, workspace_dir, parent=None):
        super().__init__(parent); self.workspace_dir = workspace_dir
        self.setWindowTitle("Project Images"); self.setFixedSize(400, 500)
        self.setStyleSheet("background:#181818;color:white;")
        layout = QVBoxLayout(self); layout.addWidget(QLabel("Project Images:"))
        self.img_list = QListWidget(); self.img_list.setSelectionMode(QAbstractItemView.MultiSelection)
        self.img_list.setStyleSheet("QListWidget{background:#0f0f12;border:2px solid #3f3f46;border-radius:8px;color:#fff;padding:8px;}QListWidget::item{padding:10px;border-radius:6px;margin-bottom:4px;}QListWidget::item:selected{background:#8b5cf6;}QListWidget::item:hover{background:#3f3f46;}")
        layout.addWidget(self.img_list)
        btns = QGridLayout()
        for txt, fn, pos in [("Table",lambda:self.insert_images("table"),(0,0)),("List",lambda:self.insert_images("list"),(0,1)),("Grid",lambda:self.insert_images("grid"),(1,0)),("Scan",self.scan_images,(1,1))]:
            b = QPushButton(txt); b.clicked.connect(fn)
            b.setStyleSheet("QPushButton{background:#27272a;color:#fff;border:2px solid #3f3f46;border-radius:8px;padding:10px 16px;font-weight:600;}QPushButton:hover{background:#3f3f46;border-color:#8b5cf6;}QPushButton:pressed{background:#8b5cf6;}")
            btns.addWidget(b, *pos)
        layout.addLayout(btns); self.scan_images()
    def scan_images(self):
        self.img_list.clear()
        for root, _, files in os.walk(self.workspace_dir):
            for f in files:
                if Path(f).suffix.lower() in {'.png','.jpg','.jpeg','.gif','.svg','.webp','.bmp'}:
                    self.img_list.addItem(os.path.relpath(os.path.join(root,f), self.workspace_dir))
    def insert_images(self, style):
        selected = [item.text() for item in self.img_list.selectedItems()]
        if not selected: return
        if style == "table":
            code = "\n| Image | Image |\n|---|---|\n"
            for i in range(0, len(selected), 2):
                img2 = f"![{selected[i+1]}]({selected[i+1]})" if i+1 < len(selected) else ""
                code += f"| ![{selected[i]}]({selected[i]}) | {img2} |\n"
        elif style == "list": code = "\n" + "\n".join([f"- ![{img}]({img})" for img in selected]) + "\n"
        else:
            code = "\n<div align='center'>\n\n"
            for img in selected: code += f"<img src='{img}' width='250' style='margin:10px;'/>\n"
            code += "\n</div>\n"
        self.insert_code.emit(code); self.accept()

class StructureDialog(QDialog):
    insert_code = Signal(str)
    def __init__(self, workspace_dir, parent=None):
        super().__init__(parent); self.workspace_dir = workspace_dir
        self.setWindowTitle("Directory Tree"); self.setFixedSize(400, 500)
        self.setStyleSheet("background:#181818;color:white;")
        layout = QVBoxLayout(self); layout.addWidget(QLabel("Directory Tree:"))
        self.tree_view = QTextEdit(); self.tree_view.setReadOnly(True)
        self.tree_view.setStyleSheet("QTextEdit{background:#0f0f12;border:2px solid #3f3f46;border-radius:8px;color:#4ade80;padding:10px;font-family:monospace;font-size:11px;}")
        layout.addWidget(self.tree_view)
        btns = QHBoxLayout()
        for txt, fn in [("Select Folder",self.select_structure_folder),("Insert",self.do_insert)]:
            b = QPushButton(txt); b.clicked.connect(fn)
            b.setStyleSheet("QPushButton{background:#27272a;color:#fff;border:2px solid #3f3f46;border-radius:8px;padding:10px 16px;font-weight:600;}QPushButton:hover{background:#3f3f46;border-color:#8b5cf6;}QPushButton:pressed{background:#8b5cf6;}")
            btns.addWidget(b)
        layout.addLayout(btns)
    def do_insert(self):
        self.insert_code.emit(f"## Project Structure\n```\n{self.tree_view.toPlainText()}\n```\n"); self.accept()
    def select_structure_folder(self):
        folder = QFileDialog.getExistingDirectory(self,"Select Folder",self.workspace_dir)
        if folder: self.tree_view.setPlainText(self.generate_tree(folder))
    def generate_tree(self, path, prefix="", max_depth=4, current_depth=0):
        if current_depth >= max_depth: return ""
        items = sorted(Path(path).iterdir(), key=lambda x:(not x.is_dir(), x.name))
        tree = ""
        for i, item in enumerate(items):
            if item.name.startswith('.') or item.name in ['node_modules','__pycache__','venv','.git']: continue
            is_last = i == len(items)-1
            tree += f"{prefix}{'└── ' if is_last else '├── '}{item.name}\n"
            if item.is_dir() and current_depth < max_depth-1:
                tree += self.generate_tree(str(item), prefix+("    " if is_last else "│   "), max_depth, current_depth+1)
        return tree

class InfoDialog(QDialog):
    insert_code = Signal(str)
    def __init__(self, workspace_dir, parent=None):
        super().__init__(parent); self.workspace_dir = workspace_dir
        self.setWindowTitle("Project Metadata"); self.setFixedSize(400, 500)
        self.setStyleSheet("background:#181818;color:white;")
        layout = QVBoxLayout(self); layout.addWidget(QLabel("Project Metadata:"))
        self.info_text = QTextEdit(); self.info_text.setReadOnly(True)
        self.info_text.setStyleSheet("QTextEdit{background:#0f0f12;border:2px solid #3f3f46;border-radius:8px;color:#58a6ff;padding:10px;font-size:12px;}")
        layout.addWidget(self.info_text)
        for txt, fn in [("Auto Detect",self.detect_project_info),("Insert",self.do_insert)]:
            b = QPushButton(txt); b.clicked.connect(fn)
            b.setStyleSheet("QPushButton{background:#27272a;color:#fff;border:2px solid #3f3f46;border-radius:8px;padding:10px 16px;font-weight:600;}QPushButton:hover{background:#3f3f46;border-color:#8b5cf6;}QPushButton:pressed{background:#8b5cf6;}")
            layout.addWidget(b)
    def do_insert(self):
        self.insert_code.emit(f"## Project Info\n{self.info_text.toPlainText()}\n"); self.accept()
    def detect_project_info(self):
        info = {}
        for fname in ["package.json","setup.py","pyproject.toml"]:
            p = Path(self.workspace_dir) / fname
            if p.exists():
                try:
                    c = p.read_text()
                    if fname == "package.json":
                        data = json.loads(c)
                        for k in ['name','description','version','author']:
                            if data.get(k): info[k.capitalize()] = data[k]
                    else:
                        if m := re.search(r'name\s*=\s*["\']([^"\']+)', c): info['Name'] = m.group(1)
                        if m := re.search(r'description\s*=\s*["\']([^"\']+)', c): info['Description'] = m.group(1)
                except: pass
        if info: self.info_text.setPlainText("\n".join([f"**{k}:** {v}" for k,v in info.items()]))
        else: self.info_text.setPlainText("No project info found\n\nSupported:\n- package.json\n- setup.py\n- pyproject.toml")


# ══════════════════════════════════════════════════════════════════════════════
# EMOJI PICKER
# ══════════════════════════════════════════════════════════════════════════════
class EmojiPicker(QListWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowFlags(Qt.ToolTip | Qt.FramelessWindowHint)
        self.setFixedSize(280, 340); self.setViewMode(QListWidget.IconMode)
        self.setGridSize(QSize(50,50)); self.setMovement(QListWidget.Static); self.setSpacing(3)
        self.setStyleSheet("QListWidget{background:#18181b;border:2px solid #3f3f46;border-radius:14px;color:#fff;outline:none;padding:12px;}QListWidget::item{border-radius:9px;padding:6px;margin:2px;}QListWidget::item:hover{background:#3f3f46;}QListWidget::item:selected{background:#8b5cf6;}")
        for e in ["🚀","✨","📝","🛠","📸","🚨","💡","📦","🔥","🎨","🔒","👤","⭐","✅","❌","🔗","📖","⚙️","📱","💻","🌐","⚡","🎉","🐛","🔑","📊","🗂","🧪","🤖","💎","🎯","🏆","🌙","☕","🦾","🔮","🌟","🧩","🎪","📌"]:
            item = QListWidgetItem(e); item.setTextAlignment(Qt.AlignCenter); item.setFont(QFont("Segoe UI Emoji",24)); self.addItem(item)
        self.hide()
    def focusOutEvent(self, e): self.hide(); super().focusOutEvent(e)
    def keyPressEvent(self, e):
        if e.key() == Qt.Key_Escape: self.hide()
        else: super().keyPressEvent(e)
    def show_at_cursor(self, editor):
        cr = editor.cursorRect()
        self.move(editor.mapToGlobal(QPoint(cr.right()+12, cr.bottom()+12))); self.show(); self.setFocus()


# ══════════════════════════════════════════════════════════════════════════════
# SUGGESTION POPUP
# ══════════════════════════════════════════════════════════════════════════════
class SuggestionPopup(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowFlags(Qt.Popup | Qt.FramelessWindowHint)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.setFixedWidth(450); self.setMaximumHeight(460)
        lay = QVBoxLayout(self); lay.setContentsMargins(0,0,0,0)
        self.list = QListWidget(); lay.addWidget(self.list)
        self.list.setStyleSheet("QListWidget{background:rgba(24,24,27,250);border:2px solid #3f3f46;border-radius:14px;color:#e4e4e7;font-size:13px;outline:none;padding:8px;}QListWidget::item{padding:12px 16px;border-radius:8px;margin-bottom:4px;}QListWidget::item:selected{background:#8b5cf6;color:#fff;}QListWidget::item:hover{background:rgba(255,255,255,0.08);}")
        self.hide()
    def showEvent(self, e):
        count = self.list.count()
        row_h = self.list.sizeHintForRow(0) if count > 0 else 32
        self.setFixedHeight(min(row_h*count+20, 460)); super().showEvent(e)
    def keyPressEvent(self, e):
        if e.key() == Qt.Key_Escape: self.hide()
        else: self.list.keyPressEvent(e)
    def clear(self): self.list.clear()
    def addItem(self, t): self.list.addItem(t)
    def currentRow(self): return self.list.currentRow()
    def setCurrentRow(self, r): self.list.setCurrentRow(r)
    def item(self, r): return self.list.item(r)
    @property
    def itemActivated(self): return self.list.itemActivated
    @property
    def itemClicked(self): return self.list.itemClicked


# ═══════════════════════════════════════════════════════════════════════════
# RECENT FILES DIALOG
# ═══════════════════════════════════════════════════════════════════════════
class RecentFilesDialog(QDialog):
    file_selected = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        t = T()
        self.setWindowTitle("Recent Files"); self.setFixedSize(580, 420)
        self.setStyleSheet(f"""
            QDialog  {{ background:{t['panel']}; color:{t['text']}; }}
            QLabel   {{ color:{t['text']}; background:transparent; }}
            QListWidget {{ background:{t['bg']}; color:{t['editor_fg']};
                           border:1px solid {t['border']}; border-radius:8px; padding:4px; }}
            QListWidget::item {{ padding:10px 14px; border-radius:6px; margin:2px; }}
            QListWidget::item:selected {{ background:{t['accent']}; color:white; }}
            QListWidget::item:hover {{ background:{t['hover']}; }}
            QPushButton {{ background:{t['hover']}; color:{t['text']};
                           border:1px solid {t['border']}; border-radius:8px;
                           padding:8px 16px; font-weight:600; font-size:12px; }}
            QPushButton:hover {{ background:{t['border']}; border-color:{t['accent']}; }}
        """)
        lay = QVBoxLayout(self); lay.setContentsMargins(20,16,20,16); lay.setSpacing(12)
        lay.addWidget(QLabel("Recent Files",
            styleSheet=f"font-size:16px;font-weight:700;color:{t['text']};background:transparent;"))

        self.list = QListWidget()
        self.list.itemDoubleClicked.connect(self._open_selected)
        lay.addWidget(self.list)

        self._populate()

        btns = QHBoxLayout(); btns.setSpacing(8)
        open_btn = QPushButton("Open"); open_btn.clicked.connect(self._open_selected); btns.addWidget(open_btn)
        clear_btn = QPushButton("Clear All"); clear_btn.clicked.connect(self._clear_all)
        clear_btn.setStyleSheet(f"background:#ef4444;color:white;border:none;border-radius:8px;padding:8px 16px;font-weight:600;")
        btns.addWidget(clear_btn)
        btns.addStretch()
        close_btn = QPushButton("Close"); close_btn.clicked.connect(self.reject); btns.addWidget(close_btn)
        lay.addLayout(btns)

    def _populate(self):
        self.list.clear()
        files = load_recent_files()
        for f in files:
            item = QListWidgetItem(f)
            exists = os.path.exists(f)
            if not exists:
                item.setForeground(QColor("#6b7280"))
                item.setToolTip("File not found")
            self.list.addItem(item)
        if not files:
            self.list.addItem("No recent files")

    def _open_selected(self, item=None):
        if item is None:
            item = self.list.currentItem()
        if item and os.path.exists(item.text()):
            self.file_selected.emit(item.text())
            self.accept()

    def _clear_all(self):
        save_recent_files([])
        self._populate()


# ═══════════════════════════════════════════════════════════════════════════
# GITHUB INTEGRATION DIALOG
# ═══════════════════════════════════════════════════════════════════════════
class GitHubWorker(QObject):
    finished = Signal(object)
    error = Signal(str)

    def __init__(self, token, method, url, body=None):
        super().__init__()
        self.token = token; self.method = method
        self.url = url; self.body = body

    def run(self):
        try:
            headers = {
                "Authorization": f"token {self.token}",
                "Accept": "application/vnd.github.v3+json",
                "User-Agent": "README-Builder-Pro"
            }
            req = urllib.request.Request(self.url, headers=headers, method=self.method)
            if self.body:
                req.add_header("Content-Type", "application/json")
                req.data = json.dumps(self.body).encode("utf-8")
            with urllib.request.urlopen(req, timeout=15) as resp:
                data = json.loads(resp.read().decode("utf-8"))
                self.finished.emit(data)
        except Exception as e:
            self.error.emit(str(e))


class GitHubDialog(QDialog):
    insert_readme = Signal(str)

    def __init__(self, markdown_text, parent=None):
        super().__init__(parent)
        self.markdown_text = markdown_text
        self._token = ""
        self._repos = []
        self._selected_repo = None
        self._threads = []
        t = T()

        self.setWindowTitle("GitHub Integration"); self.setMinimumSize(700, 560)
        self.setStyleSheet(f"""
            QDialog      {{ background:{t['panel']}; color:{t['text']}; }}
            QLabel       {{ color:{t['text']}; background:transparent; font-size:13px; }}
            QLineEdit    {{ background:{t['input_bg']}; color:{t['editor_fg']};
                           border:1px solid {t['border']}; border-radius:8px;
                           padding:8px 12px; font-size:13px; }}
            QLineEdit:focus {{ border-color:{t['accent']}; }}
            QPushButton  {{ background:{t['hover']}; color:{t['text']};
                           border:1px solid {t['border']}; border-radius:8px;
                           padding:8px 16px; font-weight:600; font-size:12px; }}
            QPushButton:hover  {{ background:{t['border']}; border-color:{t['accent']}; }}
            QPushButton#primary {{ background:{t['accent']}; border:none; color:white; }}
            QPushButton#primary:hover {{ background:{t['accent2']}; }}
            QListWidget  {{ background:{t['bg']}; color:{t['editor_fg']};
                           border:1px solid {t['border']}; border-radius:8px; padding:4px; }}
            QListWidget::item {{ padding:8px 12px; border-radius:6px; margin:2px; }}
            QListWidget::item:selected {{ background:{t['accent']}; color:white; }}
            QListWidget::item:hover {{ background:{t['hover']}; }}
            QTabWidget::pane {{ border:1px solid {t['border']}; background:{t['bg']}; border-radius:0 8px 8px 8px; }}
            QTabBar::tab {{ background:{t['panel']}; color:{t['dim']}; padding:8px 18px;
                           border:1px solid {t['border']}; border-bottom:none;
                           border-top-left-radius:6px; border-top-right-radius:6px; }}
            QTabBar::tab:selected {{ background:{t['bg']}; color:{t['text']}; }}
            QTabBar::tab:hover {{ background:{t['hover']}; color:{t['text']}; }}
            QTextEdit    {{ background:{t['bg']}; color:{t['editor_fg']};
                           border:none; padding:8px; font-size:12px; }}
            QTreeWidget  {{ background:{t['bg']}; color:{t['editor_fg']};
                           border:1px solid {t['border']}; border-radius:8px; }}
            QTreeWidget::item {{ padding:4px 8px; }}
            QTreeWidget::item:selected {{ background:{t['accent']}; }}
        """)

        main_lay = QVBoxLayout(self); main_lay.setContentsMargins(20,16,20,16); main_lay.setSpacing(12)

        # Title
        main_lay.addWidget(QLabel("GitHub Integration",
            styleSheet=f"font-size:18px;font-weight:700;color:{t['text']};background:transparent;"))

        # PAT Token row
        token_row = QHBoxLayout(); token_row.setSpacing(8)
        token_row.addWidget(QLabel("PAT Token:"))
        self.token_input = QLineEdit()
        self.token_input.setPlaceholderText("ghp_xxxxxxxxxxxxxxxxxxxx")
        self.token_input.setEchoMode(QLineEdit.Password)
        self._load_token()
        token_row.addWidget(self.token_input)
        self.connect_btn = QPushButton("Connect")
        self.connect_btn.setObjectName("primary")
        self.connect_btn.clicked.connect(self.connect_github)
        token_row.addWidget(self.connect_btn)
        show_btn = QPushButton("Show")
        show_btn.setFixedWidth(60)
        show_btn.clicked.connect(lambda: self.token_input.setEchoMode(
            QLineEdit.Normal if self.token_input.echoMode() == QLineEdit.Password else QLineEdit.Password))
        token_row.addWidget(show_btn)
        main_lay.addLayout(token_row)

        self.status_lbl = QLabel("")
        self.status_lbl.setStyleSheet(f"color:{t['dim']};font-size:11px;background:transparent;")
        main_lay.addWidget(self.status_lbl)

        # Repo list + tabs
        content_split = QSplitter(Qt.Horizontal)

        # Repos list
        repo_widget = QWidget(); repo_lay = QVBoxLayout(repo_widget)
        repo_lay.setContentsMargins(0,0,0,0); repo_lay.setSpacing(6)
        repo_lay.addWidget(QLabel("Repositories:"))
        self.repo_search = QLineEdit(); self.repo_search.setPlaceholderText("Search repos...")
        self.repo_search.textChanged.connect(self._filter_repos)
        repo_lay.addWidget(self.repo_search)
        self.repo_list = QListWidget()
        self.repo_list.itemClicked.connect(self._on_repo_selected)
        repo_lay.addWidget(self.repo_list)
        content_split.addWidget(repo_widget)

        # Tabs
        self.tabs = QTabWidget()

        # README tab
        readme_tab = QWidget(); readme_lay = QVBoxLayout(readme_tab); readme_lay.setContentsMargins(8,8,8,8)
        readme_btn_row = QHBoxLayout()
        self.import_btn = QPushButton("Import README"); self.import_btn.clicked.connect(self.import_readme)
        self.upload_btn = QPushButton("Upload README"); self.upload_btn.setObjectName("primary")
        self.upload_btn.clicked.connect(self.upload_readme)
        readme_btn_row.addWidget(self.import_btn); readme_btn_row.addWidget(self.upload_btn)
        readme_lay.addLayout(readme_btn_row)
        # File upload options
        file_row = QHBoxLayout()
        file_row.addWidget(QLabel("Also upload:"))
        self.upload_file_btn = QPushButton("Add File/Folder")
        self.upload_file_btn.clicked.connect(self.add_upload_file)
        file_row.addWidget(self.upload_file_btn)
        readme_lay.addLayout(file_row)
        self.upload_files_list = QListWidget(); self.upload_files_list.setFixedHeight(80)
        readme_lay.addWidget(self.upload_files_list)
        self.commit_msg = QLineEdit(); self.commit_msg.setPlaceholderText("Commit message (optional)")
        readme_lay.addWidget(self.commit_msg)
        self.tabs.addTab(readme_tab, "README")

        # Issues tab
        issues_tab = QWidget(); issues_lay = QVBoxLayout(issues_tab); issues_lay.setContentsMargins(8,8,8,8)
        issues_lay.addWidget(QLabel("Issues & Pull Requests:"))
        self.issues_tree = QTreeWidget()
        self.issues_tree.setHeaderLabels(["#", "Title", "State", "Type"])
        self.issues_tree.setColumnWidth(0, 50); self.issues_tree.setColumnWidth(1, 280)
        self.issues_tree.setColumnWidth(2, 60)
        issues_lay.addWidget(self.issues_tree)
        refresh_issues_btn = QPushButton("Refresh Issues/PRs")
        refresh_issues_btn.clicked.connect(self.fetch_issues)
        issues_lay.addWidget(refresh_issues_btn)
        self.tabs.addTab(issues_tab, "Issues & PRs")

        # Contributors tab
        contrib_tab = QWidget(); contrib_lay = QVBoxLayout(contrib_tab); contrib_lay.setContentsMargins(8,8,8,8)
        contrib_lay.addWidget(QLabel("Contributors:"))
        self.contrib_list = QListWidget()
        contrib_lay.addWidget(self.contrib_list)
        contrib_btn_row = QHBoxLayout()
        refresh_contrib_btn = QPushButton("Refresh"); refresh_contrib_btn.clicked.connect(self.fetch_contributors); contrib_btn_row.addWidget(refresh_contrib_btn)
        insert_contrib_btn = QPushButton("Insert into README"); insert_contrib_btn.clicked.connect(self.insert_contributors_section); contrib_btn_row.addWidget(insert_contrib_btn)
        contrib_lay.addLayout(contrib_btn_row)
        self.tabs.addTab(contrib_tab, "Contributors")

        content_split.addWidget(self.tabs)
        content_split.setSizes([220, 460])
        main_lay.addWidget(content_split)

    def _load_token(self):
        p = os.path.expanduser("~/.readme_builder_gh_token")
        if os.path.exists(p):
            try:
                with open(p,"r") as f: self.token_input.setText(f.read().strip())
            except: pass

    def _save_token(self, token):
        p = os.path.expanduser("~/.readme_builder_gh_token")
        try:
            with open(p,"w") as f: f.write(token)
        except: pass

    def _set_status(self, msg, color="#8b949e"):
        self.status_lbl.setText(msg)
        self.status_lbl.setStyleSheet(f"color:{color};font-size:11px;background:transparent;")

    def _gh_request(self, method, url, callback, body=None):
        token = self.token_input.text().strip()
        if not token: self._set_status("Please enter a PAT token", "#f87171"); return
        worker = GitHubWorker(token, method, url, body)
        thread = QThread()
        worker.moveToThread(thread)
        worker.finished.connect(callback)
        worker.error.connect(lambda e: self._set_status(f"Error: {e}", "#f87171"))
        thread.started.connect(worker.run)
        thread.start()
        self._threads.append((thread, worker))

    def connect_github(self):
        token = self.token_input.text().strip()
        if not token: self._set_status("Enter a PAT token first", "#f87171"); return
        self._token = token
        self._save_token(token)
        self._set_status("Connecting...", "#fbbf24")
        self._gh_request("GET", "https://api.github.com/user/repos?per_page=100&sort=updated",
                         self._on_repos_loaded)

    def _on_repos_loaded(self, data):
        self._repos = data if isinstance(data, list) else []
        self.repo_list.clear()
        for r in self._repos:
            item = QListWidgetItem(r.get("full_name",""))
            item.setData(Qt.UserRole, r)
            self.repo_list.addItem(item)
        count = len(self._repos)
        self._set_status(f"Connected — {count} repositories loaded", "#4ade80")

    def _filter_repos(self, text):
        for i in range(self.repo_list.count()):
            item = self.repo_list.item(i)
            item.setHidden(text.lower() not in item.text().lower())

    def _on_repo_selected(self, item):
        self._selected_repo = item.data(Qt.UserRole)

    def import_readme(self):
        if not self._selected_repo: self._set_status("Select a repository first", "#f87171"); return
        full = self._selected_repo.get("full_name","")
        for branch in ["main", "master"]:
            url = f"https://raw.githubusercontent.com/{full}/{branch}/README.md"
            try:
                req = urllib.request.Request(url, headers={"User-Agent":"README-Builder-Pro"})
                with urllib.request.urlopen(req, timeout=10) as resp:
                    content = resp.read().decode("utf-8")
                    self.insert_readme.emit(content)
                    self._set_status(f"Imported README from {full} ({branch})", "#4ade80")
                    self.accept(); return
            except: continue
        self._set_status("README not found in main or master branch", "#f87171")

    def upload_readme(self):
        if not self._selected_repo: self._set_status("Select a repository first", "#f87171"); return
        full = self._selected_repo.get("full_name","")
        token = self.token_input.text().strip()
        if not token: return

        import base64
        content_b64 = base64.b64encode(self.markdown_text.encode("utf-8")).decode("utf-8")
        commit_msg = self.commit_msg.text().strip() or "Update README.md via README Builder Pro"

        # Check if README exists to get SHA
        self._set_status("Uploading README...", "#fbbf24")
        sha = self._get_file_sha(full, "README.md", token)
        body = {"message": commit_msg, "content": content_b64}
        if sha: body["sha"] = sha

        url = f"https://api.github.com/repos/{full}/contents/README.md"
        self._gh_request("PUT", url, lambda _: self._on_readme_uploaded(full), body=body)

        # Upload extra files
        for i in range(self.upload_files_list.count()):
            fpath = self.upload_files_list.item(i).text()
            if os.path.isfile(fpath):
                self._upload_single_file(full, fpath, token, commit_msg)

    def _get_file_sha(self, full_name, path, token):
        try:
            url = f"https://api.github.com/repos/{full_name}/contents/{path}"
            headers = {"Authorization": f"token {token}", "User-Agent":"README-Builder-Pro"}
            req = urllib.request.Request(url, headers=headers)
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = json.loads(resp.read().decode("utf-8"))
                return data.get("sha")
        except:
            return None

    def _upload_single_file(self, full_name, file_path, token, commit_msg):
        import base64
        fname = os.path.basename(file_path)
        try:
            with open(file_path, "rb") as f:
                content = base64.b64encode(f.read()).decode("utf-8")
            sha = self._get_file_sha(full_name, fname, token)
            body = {"message": f"Upload {fname} via README Builder Pro", "content": content}
            if sha: body["sha"] = sha
            url = f"https://api.github.com/repos/{full_name}/contents/{fname}"
            self._gh_request("PUT", url, lambda _: None, body=body)
        except Exception as e:
            self._set_status(f"Error uploading {fname}: {e}", "#f87171")

    def _on_readme_uploaded(self, full_name):
        self._set_status(f"README uploaded to {full_name} successfully!", "#4ade80")

    def add_upload_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select File to Upload", "", "All Files (*.*)")
        if path:
            self.upload_files_list.addItem(path)

    def fetch_issues(self):
        if not self._selected_repo: self._set_status("Select a repository first", "#f87171"); return
        full = self._selected_repo.get("full_name","")
        url = f"https://api.github.com/repos/{full}/issues?state=open&per_page=50"
        self._set_status("Loading issues...", "#fbbf24")
        self._gh_request("GET", url, self._on_issues_loaded)

    def _on_issues_loaded(self, data):
        self.issues_tree.clear()
        if not isinstance(data, list): return
        for issue in data:
            num = str(issue.get("number",""))
            title = issue.get("title","")
            state = issue.get("state","")
            is_pr = "pull_request" in issue
            itype = "PR" if is_pr else "Issue"
            item = QTreeWidgetItem([num, title, state, itype])
            if is_pr: item.setForeground(1, QColor("#8b5cf6"))
            self.issues_tree.addTopLevelItem(item)
        self._set_status(f"Loaded {len(data)} items", "#4ade80")

    def fetch_contributors(self):
        if not self._selected_repo: self._set_status("Select a repository first", "#f87171"); return
        full = self._selected_repo.get("full_name","")
        url = f"https://api.github.com/repos/{full}/contributors?per_page=50"
        self._set_status("Loading contributors...", "#fbbf24")
        self._gh_request("GET", url, self._on_contributors_loaded)

    def _on_contributors_loaded(self, data):
        self.contrib_list.clear()
        if not isinstance(data, list): return
        for c in data:
            login = c.get("login","")
            contributions = c.get("contributions", 0)
            item = QListWidgetItem(f"{login}  ({contributions} commits)")
            item.setData(Qt.UserRole, c)
            self.contrib_list.addItem(item)
        self._set_status(f"Loaded {len(data)} contributors", "#4ade80")

    def insert_contributors_section(self):
        items = []
        for i in range(self.contrib_list.count()):
            item = self.contrib_list.item(i)
            data = item.data(Qt.UserRole)
            if data:
                login = data.get("login","")
                url = data.get("html_url","")
                contrib = data.get("contributions", 0)
                items.append(f"| [@{login}]({url}) | {contrib} |")
        if not items:
            self._set_status("No contributors to insert", "#f87171"); return
        section = "## Contributors\n\n| Contributor | Commits |\n|------------|--------|\n"
        section += "\n".join(items) + "\n"
        self.insert_readme.emit(section)
        self._set_status("Contributors section inserted!", "#4ade80")
        self.accept()


# ═══════════════════════════════════════════════════════════════════════════
# OFFLINE TRANSLATION DIALOG
# ═══════════════════════════════════════════════════════════════════════════
class TranslateDialog(QDialog):
    translated = Signal(str)

    def __init__(self, markdown_text, parent=None):
        super().__init__(parent)
        self.markdown_text = markdown_text
        t = T()
        self.setWindowTitle("Offline Translation"); self.setFixedSize(520, 380)
        self.setStyleSheet(f"""
            QDialog   {{ background:{t['panel']}; color:{t['text']}; }}
            QLabel    {{ color:{t['text']}; background:transparent; font-size:13px; }}
            QComboBox {{ background:{t['input_bg']}; color:{t['editor_fg']};
                        border:1px solid {t['border']}; border-radius:8px;
                        padding:8px 12px; font-size:13px; }}
            QComboBox QAbstractItemView {{ background:{t['panel']}; color:{t['text']}; border:1px solid {t['border']}; }}
            QPushButton {{ background:{t['hover']}; color:{t['text']};
                          border:1px solid {t['border']}; border-radius:8px;
                          padding:8px 16px; font-weight:600; font-size:12px; }}
            QPushButton:hover {{ background:{t['border']}; border-color:{t['accent']}; }}
            QPushButton#primary {{ background:{t['accent']}; border:none; color:white; }}
            QPushButton#primary:hover {{ background:{t['accent2']}; }}
            QTextEdit {{ background:{t['bg']}; color:{t['editor_fg']};
                        border:1px solid {t['border']}; border-radius:8px; padding:8px; }}
        """)
        lay = QVBoxLayout(self); lay.setContentsMargins(20,16,20,16); lay.setSpacing(12)
        lay.addWidget(QLabel("Offline Markdown Translation",
            styleSheet=f"font-size:16px;font-weight:700;color:{t['text']};background:transparent;"))
        lay.addWidget(QLabel(
            "Translates section headings and standard terms offline — no internet or AI required.",
            styleSheet=f"color:{t['dim']};font-size:12px;background:transparent;"))

        lang_row = QHBoxLayout(); lang_row.setSpacing(10)
        lang_row.addWidget(QLabel("Target Language:"))
        self.lang_combo = QComboBox()
        for code, name in LANG_NAMES.items():
            self.lang_combo.addItem(name, code)
        lang_row.addWidget(self.lang_combo)
        lang_row.addStretch()
        lay.addLayout(lang_row)

        lay.addWidget(QLabel("Preview (headings that will be translated):"))
        self.preview = QTextEdit(); self.preview.setReadOnly(True); self.preview.setFixedHeight(160)
        # Show headings
        headings = [l for l in markdown_text.split("\n") if l.startswith("#")]
        self.preview.setPlainText("\n".join(headings[:20]))
        lay.addWidget(self.preview)

        self.lang_combo.currentIndexChanged.connect(self._update_preview)
        self._update_preview()

        btns = QHBoxLayout(); btns.setSpacing(8)
        tr_btn = QPushButton("Translate"); tr_btn.setObjectName("primary")
        tr_btn.clicked.connect(self._do_translate); btns.addWidget(tr_btn)
        cancel_btn = QPushButton("Cancel"); cancel_btn.clicked.connect(self.reject); btns.addWidget(cancel_btn)
        lay.addLayout(btns)

    def _update_preview(self):
        code = self.lang_combo.currentData()
        headings = [l for l in self.markdown_text.split("\n") if l.startswith("#")]
        translated_lines = []
        for h in headings[:20]:
            result = translate_markdown_offline(h, code)
            result = result.replace("<!-- dir: rtl -->\n", "")
            translated_lines.append(result)
        self.preview.setPlainText("\n".join(translated_lines))

    def _do_translate(self):
        code = self.lang_combo.currentData()
        result = translate_markdown_offline(self.markdown_text, code)
        self.translated.emit(result)
        self.accept()


# ═══════════════════════════════════════════════════════════════════════════
# SMART EDITOR  (line numbers + highlighter + Undo/Redo history + context menu)
# ═══════════════════════════════════════════════════════════════════════════
class SmartEditor(QTextEdit):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.popup = SuggestionPopup(self)
        self.popup.itemClicked.connect(self.insert_suggestion)
        self.popup.itemActivated.connect(self.insert_suggestion)
        self.data = dict(TAB_SNIPPETS)
        self.load_snippets(); self.refresh_popup_items()
        self.highlighter = MarkdownHighlighter(self.document())
        self.line_number_area = LineNumberArea(self)
        self.document().blockCountChanged.connect(self.update_line_number_area_width)
        self.verticalScrollBar().valueChanged.connect(self.update_line_number_area)
        self.cursorPositionChanged.connect(self._highlight_cur_line)
        self.update_line_number_area_width()

        # Enhanced Undo/Redo
        self._undo_stack = QUndoStack(self)
        self._undo_stack.setUndoLimit(500)
        self._last_text = ""
        self._snapshot_text = ""
        self._undo_redo_active = False
        self._undo_timer = QTimer()
        self._undo_timer.setSingleShot(True)
        self._undo_timer.setInterval(600)
        self._undo_timer.timeout.connect(self._commit_undo)
        self.textChanged.connect(self._on_text_changed)

    def _on_text_changed(self):
        if self._undo_redo_active: return
        if not self._undo_timer.isActive():
            self._snapshot_text = self._last_text
        self._undo_timer.start()

    def _commit_undo(self):
        if self._undo_redo_active: return
        new_text = self.toPlainText()
        if new_text != self._snapshot_text:
            cmd = TextChangeCommand(self, self._snapshot_text, new_text,
                                    self.textCursor().position())
            self._undo_stack.push(cmd)
            self._last_text = new_text
            self._snapshot_text = new_text

    def undo(self):
        if self._undo_timer.isActive():
            self._undo_timer.stop()
            self._commit_undo()
        if self._undo_stack.canUndo():
            self._undo_stack.undo()
        else:
            super().undo()

    def redo(self):
        if self._undo_stack.canRedo():
            self._undo_stack.redo()
        else:
            super().redo()


    # ── line numbers ─────────────────────────────────────────────────────
    def line_number_area_width(self):
        return 14 + self.fontMetrics().horizontalAdvance("9") * max(len(str(self.document().blockCount())), 3)

    def update_line_number_area_width(self):
        self.setViewportMargins(self.line_number_area_width(), 0, 0, 0)

    def update_line_number_area(self):
        self.line_number_area.update()
        self.line_number_area.resize(self.line_number_area_width(), self.height())

    def resizeEvent(self, event):
        super().resizeEvent(event)
        cr = self.contentsRect()
        self.line_number_area.setGeometry(
            QRect(cr.left(), cr.top(), self.line_number_area_width(), cr.height()))

    def _highlight_cur_line(self):
        t = T(); sel = QTextEdit.ExtraSelection()
        sel.format.setBackground(QColor(t["hover"]))
        sel.format.setProperty(QTextCharFormat.FullWidthSelection, True)
        sel.cursor = self.textCursor(); sel.cursor.clearSelection()
        self.setExtraSelections([sel])

    def line_number_area_paint_event(self, event):
        t = T(); painter = QPainter(self.line_number_area)
        painter.fillRect(event.rect(), QColor(t["input_bg"]))
        block = self.document().begin(); bn = 0
        cur_bn = self.textCursor().blockNumber()
        while block.isValid():
            y = int(self.document().documentLayout().blockBoundingRect(block).top()) \
                - self.verticalScrollBar().value() + self.contentsMargins().top()
            if y > event.rect().bottom(): break
            if y + self.fontMetrics().height() >= event.rect().top():
                painter.setPen(QColor(t["accent"] if bn == cur_bn else t["border"]))
                painter.setFont(self.font())
                painter.drawText(0, y, self.line_number_area.width()-6,
                                 self.fontMetrics().height(), Qt.AlignRight, str(bn+1))
            block = block.next(); bn += 1

    # ── snippets ─────────────────────────────────────────────────────────
    def load_snippets(self):
        if os.path.exists(SNIPPETS_FILE):
            try:
                with open(SNIPPETS_FILE,"r",encoding="utf-8") as f:
                    self.data.update(json.load(f))
            except: pass

    def save_snippets(self):
        try:
            custom = {k:v for k,v in self.data.items() if k.startswith("Custom ")}
            with open(SNIPPETS_FILE,"w",encoding="utf-8") as f:
                json.dump(custom, f, ensure_ascii=False, indent=4)
        except: pass

    def refresh_popup_items(self):
        self.popup.clear()
        for k in self.data: self.popup.addItem(k)

    def save_custom_snippet(self):
        sel = self.textCursor().selectedText().replace("\u2029","\n")
        if not sel.strip(): QMessageBox.warning(self,"Warning","Select text first!"); return
        name, ok = ModernInputDialog.get_text(self,"Save Snippet","Snippet name:","My Snippet")
        if ok and name:
            self.data[f"Custom {name}"] = sel; self.save_snippets()
            self.refresh_popup_items(); QMessageBox.information(self,"Done","Snippet saved!")

    def insert_markdown(self, prefix, suffix=""):
        cur = self.textCursor()
        if cur.hasSelection(): cur.insertText(f"{prefix}{cur.selectedText()}{suffix}")
        else:
            cur.insertText(f"{prefix}{suffix}")
            if suffix:
                cur.setPosition(cur.position()-len(suffix)); self.setTextCursor(cur)

    # ── themed context menu ───────────────────────────────────────────────
    def contextMenuEvent(self, event):
        t = T()
        ms = f"""
            QMenu {{
                background:{t['panel']}; border:1px solid {t['border']};
                border-radius:10px; padding:6px 4px;
                color:{t['text']}; font-size:12px; font-weight:500;
            }}
            QMenu::item {{ padding:8px 28px 8px 20px; border-radius:6px; margin:2px 4px; }}
            QMenu::item:selected {{ background:{t['accent']}; color:#fff; }}
            QMenu::item:disabled {{ color:{t['dim']}; }}
            QMenu::separator {{ height:1px; background:{t['border']}; margin:4px 10px; }}
            QMenu::right-arrow {{ color:{t['dim']}; }}
        """
        m = QMenu(self); m.setStyleSheet(ms)
        ua=m.addAction("Undo"); ua.setShortcut("Ctrl+Z")
        ua.setEnabled(self._undo_stack.canUndo()); ua.triggered.connect(self.undo)
        ra=m.addAction("Redo"); ra.setShortcut("Ctrl+Y")
        ra.setEnabled(self._undo_stack.canRedo()); ra.triggered.connect(self.redo)
        m.addSeparator()
        ca=m.addAction("Cut"); ca.setEnabled(self.textCursor().hasSelection()); ca.triggered.connect(self.cut)
        cp=m.addAction("Copy"); cp.setEnabled(self.textCursor().hasSelection()); cp.triggered.connect(self.copy)
        pa=m.addAction("Paste"); pa.triggered.connect(self.paste)
        sa=m.addAction("Select All"); sa.triggered.connect(self.selectAll)
        m.addSeparator()

        def sub(label):
            s = QMenu(label, m); s.setStyleSheet(ms); return s

        hm = sub("Headers")
        for i in range(1,5):
            hm.addAction(f"H{i}  {'#'*i} Heading").triggered.connect(
                lambda _, n=i: self.insert_markdown("#"*n+" ",""))
        m.addMenu(hm)

        fm = sub("Format")
        for lbl, pre, suf in [
            ("Bold **text**","**","**"),("Italic *text*","*","*"),
            ("Inline Code `code`","`","`"),("Code Block","```\n","\n```"),
            ("Strikethrough","~~","~~"),
        ]:
            fm.addAction(f"{lbl}").triggered.connect(
                lambda _, p=pre, s=suf: self.insert_markdown(p,s))
        m.addMenu(fm)

        tm = sub("Tables")
        for lbl, txt in [
            ("2x2", "\n| Col 1 | Col 2 |\n|-------|-------|\n| Val 1 | Val 2 |\n"),
            ("3x3", "\n| A | B | C |\n|---|---|---|\n| 1 | 2 | 3 |\n"),
            ("API", "\n| Method | Endpoint | Desc |\n|--------|----------|------|\n| `GET` | `/items` | List |\n"),
        ]:
            tm.addAction(f"{lbl}").triggered.connect(lambda _, x=txt: self.insertPlainText(x))
        m.addMenu(tm)

        bm = sub("Badges")
        for lbl, txt in [
            ("Build", "[![Build](https://img.shields.io/badge/build-passing-brightgreen?style=flat-square)](#)\n"),
            ("MIT",   "[![License](https://img.shields.io/badge/license-MIT-blue?style=flat-square)](#)\n"),
            ("Python","[![Python](https://img.shields.io/badge/Python-3.8+-blue?style=flat-square&logo=python)](#)\n"),
            ("v1.0",  "[![Version](https://img.shields.io/badge/version-1.0.0-purple?style=flat-square)](#)\n"),
        ]:
            bm.addAction(f"{lbl}").triggered.connect(lambda _, x=txt: self.insertPlainText(x))
        m.addMenu(bm)

        am = sub("Alerts")
        for lbl, txt in [
            ("Note",     "> [!NOTE]\n> Your note here.\n"),
            ("Warning",  "> [!WARNING]\n> Your warning.\n"),
            ("Tip",      "> [!TIP]\n> Your tip.\n"),
            ("Important","> [!IMPORTANT]\n> Important info.\n"),
            ("Caution",  "> [!CAUTION]\n> Caution!\n"),
        ]:
            am.addAction(f"{lbl}").triggered.connect(lambda _, x=txt: self.insertPlainText(x))
        m.addMenu(am)

        if self.data:
            m.addSeparator()
            sm = sub("My Snippets")
            for key in sorted(self.data):
                sm.addAction(f"{key}").triggered.connect(
                    lambda _, k=key: self.insertPlainText(self.data[k]))
            m.addMenu(sm)

        m.exec(event.globalPos())

    def keyPressEvent(self, e):
        if e.key() == Qt.Key_Tab:
            cr = self.cursorRect()
            self.popup.move(self.mapToGlobal(QPoint(cr.right()+16, cr.bottom()+12)))
            self.popup.show(); self.popup.setFocus(); self.popup.setCurrentRow(0)
        elif self.popup.isVisible() and e.key() in (Qt.Key_Enter, Qt.Key_Return):
            self.insert_suggestion(self.popup.item(self.popup.currentRow()))
        elif self.popup.isVisible() and e.key() == Qt.Key_Escape:
            self.popup.hide(); self.setFocus()
        elif e.key() == Qt.Key_Z and e.modifiers() & Qt.ControlModifier:
            self.undo()
        elif e.key() == Qt.Key_Y and e.modifiers() & Qt.ControlModifier:
            self.redo()
        elif e.key() == Qt.Key_V and e.modifiers() & Qt.ControlModifier:
            self._smart_paste()
        else: super().keyPressEvent(e)

    def _smart_paste(self):
        """Smart Paste: auto-formats URLs and image links."""
        import re as _re
        from PySide6.QtWidgets import QApplication as _QApp
        clipboard = _QApp.clipboard()
        clip_text = clipboard.text().strip()
        cur = self.textCursor()
        selected = cur.selectedText().strip()

        # Detect image URL
        is_image_url = bool(_re.match(
            r'https?://\S+\.(?:png|jpe?g|gif|webp|svg|bmp)(\?\S*)?$',
            clip_text, _re.IGNORECASE))
        # Detect any URL
        is_url = bool(_re.match(r'https?://\S+', clip_text))

        if selected and is_image_url:
            cur.insertText(f"![{selected}]({clip_text})")
            self.setTextCursor(cur)
        elif selected and is_url:
            cur.insertText(f"[{selected}]({clip_text})")
            self.setTextCursor(cur)
        elif not selected and is_image_url:
            alt = clip_text.rstrip('/').split('/')[-1].split('.')[0]
            cur.insertText(f"![{alt}]({clip_text})")
            self.setTextCursor(cur)
        else:
            self.paste()

    def insert_suggestion(self, item):
        if not item: return
        self.insertPlainText(self.data[item.text()]); self.popup.hide(); self.setFocus()


# ═══════════════════════════════════════════════════════════════════════════
# MARKDOWN TABLE EDITOR
# ═══════════════════════════════════════════════════════════════════════════
class TableEditorDialog(QDialog):
    insert_code = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        t = T()
        self.setWindowTitle("Markdown Table Editor")
        self.setFixedSize(700, 480)
        self.setStyleSheet(f"""
            QDialog    {{ background:{t['panel']}; color:{t['text']}; }}
            QLabel     {{ color:{t['text']}; background:transparent; font-size:13px; }}
            QSpinBox   {{ background:{t['input_bg']}; color:{t['editor_fg']};
                          border:1px solid {t['border']}; border-radius:6px; padding:4px 8px; }}
            QPushButton {{ background:{t['hover']}; color:{t['text']};
                           border:1px solid {t['border']}; border-radius:8px;
                           padding:8px 14px; font-size:12px; font-weight:600; }}
            QPushButton:hover  {{ background:{t['border']}; border-color:{t['accent']}; }}
            QPushButton#primary {{ background:{t['accent']}; border:none; color:white; }}
            QPushButton#primary:hover {{ background:{t['accent2']}; }}
            QTextEdit  {{ background:{t['bg']}; color:{t['editor_fg']};
                          border:1px solid {t['border']}; border-radius:8px;
                          padding:8px; font-family:monospace; font-size:12px; }}
        """)

        lay = QVBoxLayout(self); lay.setContentsMargins(20,16,20,16); lay.setSpacing(12)
        lay.addWidget(QLabel("Markdown Table Editor",
            styleSheet=f"font-size:16px;font-weight:700;color:{t['text']};background:transparent;"))

        # Rows / Cols controls
        ctrl_row = QHBoxLayout(); ctrl_row.setSpacing(12)
        ctrl_row.addWidget(QLabel("Rows:"))
        self.rows_spin = QSpinBox(); self.rows_spin.setRange(1, 20); self.rows_spin.setValue(3)
        ctrl_row.addWidget(self.rows_spin)
        ctrl_row.addWidget(QLabel("Columns:"))
        self.cols_spin = QSpinBox(); self.cols_spin.setRange(1, 10); self.cols_spin.setValue(3)
        ctrl_row.addWidget(self.cols_spin)
        build_btn = QPushButton("Build Grid"); build_btn.clicked.connect(self.build_grid)
        ctrl_row.addWidget(build_btn)
        ctrl_row.addStretch()
        # Alignment
        ctrl_row.addWidget(QLabel("Align:"))
        self.align_combo = QComboBox()
        self.align_combo.addItems(["Left  :---", "Center :---:", "Right ---:"])
        self.align_combo.setStyleSheet(f"""
            QComboBox {{ background:{t['input_bg']}; color:{t['editor_fg']};
                        border:1px solid {t['border']}; border-radius:6px; padding:4px 8px; }}
            QComboBox QAbstractItemView {{ background:{t['panel']}; color:{t['text']};
                                           border:1px solid {t['border']}; }}
        """)
        ctrl_row.addWidget(self.align_combo)
        lay.addLayout(ctrl_row)

        # Table grid area
        self.grid_scroll = QScrollArea(); self.grid_scroll.setWidgetResizable(True)
        self.grid_scroll.setStyleSheet(f"background:{t['bg']};border:1px solid {t['border']};border-radius:8px;")
        self.grid_widget = QWidget(); self.grid_widget.setStyleSheet(f"background:{t['bg']};")
        self.grid_layout = QGridLayout(self.grid_widget); self.grid_layout.setSpacing(3)
        self.grid_scroll.setWidget(self.grid_widget)
        lay.addWidget(self.grid_scroll)
        self.cells = []  # list of lists of QLineEdit

        # Preview
        lay.addWidget(QLabel("Markdown Preview:"))
        self.preview_edit = QTextEdit(); self.preview_edit.setReadOnly(True); self.preview_edit.setFixedHeight(80)
        lay.addWidget(self.preview_edit)

        btns = QHBoxLayout(); btns.setSpacing(8)
        preview_btn = QPushButton("Preview"); preview_btn.clicked.connect(self.update_preview)
        btns.addWidget(preview_btn)
        btns.addStretch()
        cancel_btn = QPushButton("Cancel"); cancel_btn.clicked.connect(self.reject)
        btns.addWidget(cancel_btn)
        insert_btn = QPushButton("Insert Table"); insert_btn.setObjectName("primary")
        insert_btn.clicked.connect(self._insert); btns.addWidget(insert_btn)
        lay.addLayout(btns)

        self.build_grid()

    def _cell_style(self):
        t = T()
        return (f"background:{t['input_bg']};color:{t['editor_fg']};"
                f"border:1px solid {t['border']};border-radius:4px;padding:4px 6px;font-size:12px;")

    def build_grid(self):
        # Clear old cells
        for r in self.cells:
            for c in r: c.deleteLater()
        self.cells = []
        while self.grid_layout.count():
            item = self.grid_layout.takeAt(0)
            if item.widget(): item.widget().deleteLater()

        rows = self.rows_spin.value(); cols = self.cols_spin.value()
        t = T()
        for r in range(rows):
            row_cells = []
            for c in range(cols):
                inp = QLineEdit()
                inp.setPlaceholderText("Header" if r == 0 else f"R{r}C{c+1}")
                inp.setStyleSheet(self._cell_style())
                inp.setMinimumWidth(80)
                if r == 0:
                    inp.setStyleSheet(self._cell_style() +
                        f"font-weight:700;border-color:{t['accent']};")
                inp.textChanged.connect(self.update_preview)
                self.grid_layout.addWidget(inp, r, c)
                row_cells.append(inp)
            self.cells.append(row_cells)
        self.update_preview()

    def _get_align_sep(self, cols):
        a = self.align_combo.currentIndex()
        if a == 0: sep = ":---"
        elif a == 1: sep = ":---:"
        else: sep = "---:"
        return "| " + " | ".join([sep] * cols) + " |"

    def generate_markdown(self) -> str:
        if not self.cells: return ""
        cols = len(self.cells[0])
        lines = []
        for r, row in enumerate(self.cells):
            vals = [c.text().strip() or (f"Header {i+1}" if r==0 else " ") for i, c in enumerate(row)]
            lines.append("| " + " | ".join(vals) + " |")
            if r == 0:
                lines.append(self._get_align_sep(cols))
        return "\n" + "\n".join(lines) + "\n"

    def update_preview(self):
        self.preview_edit.setPlainText(self.generate_markdown())

    def _insert(self):
        md = self.generate_markdown()
        if md.strip():
            self.insert_code.emit(md); self.accept()


# ═══════════════════════════════════════════════════════════════════════════
# QUICK PROJECT DIALOG
# ═══════════════════════════════════════════════════════════════════════════
class QuickProjectDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        t = T()
        self.setWindowTitle("Quick Templates"); self.setFixedSize(640, 500)
        self.setStyleSheet(f"""
            QDialog    {{ background:{t['panel']}; color:{t['text']}; }}
            QLabel     {{ color:{t['text']}; background:transparent; }}
            QLineEdit  {{ background:{t['input_bg']}; color:{t['editor_fg']};
                          border:2px solid {t['border']}; border-radius:8px; padding:8px 12px; font-size:13px; }}
            QLineEdit:focus {{ border-color:{t['accent']}; }}
            QPushButton {{ background:{t['hover']}; color:{t['text']};
                           border:1px solid {t['border']}; border-radius:8px;
                           padding:8px 14px; font-size:12px; font-weight:600; }}
            QPushButton:hover  {{ background:{t['border']}; border-color:{t['accent']}; color:white; }}
            QPushButton#apply  {{ background:{t['accent']}; border:none; color:white; }}
            QPushButton#apply:hover {{ background:{t['accent2']}; }}
        """)
        lay = QVBoxLayout(self); lay.setContentsMargins(24,20,24,20); lay.setSpacing(14)
        lay.addWidget(QLabel("Quick Templates",
            styleSheet=f"color:{t['text']};font-size:18px;font-weight:600;background:transparent;"))
        lay.addWidget(QLabel("Select a template and enter your project name:",
            styleSheet=f"color:{t['dim']};font-size:12px;background:transparent;"))
        row = QHBoxLayout()
        row.addWidget(QLabel("Project Name:",
            styleSheet=f"color:{t['dim']};font-size:12px;min-width:100px;background:transparent;"))
        self.name_input = QLineEdit(); self.name_input.setPlaceholderText("My Awesome Project")
        row.addWidget(self.name_input); lay.addLayout(row)

        scroll = QScrollArea(); scroll.setWidgetResizable(True)
        scroll.setStyleSheet(f"QScrollArea{{border:none;background:{t['bg']};}}")
        inner = QWidget(); inner.setStyleSheet(f"background:{t['bg']};")
        grid = QGridLayout(inner); grid.setSpacing(8); grid.setContentsMargins(4,4,4,4)
        self.selected_key = None; self.selected_template = None; self._btns = []
        for i, key in enumerate(QUICK_TEMPLATES):
            btn = QPushButton(key); btn.setFixedHeight(46)
            btn.setStyleSheet(f"""QPushButton{{background:{t['hover']};color:{t['text']};
                border:1px solid {t['border']};border-radius:8px;font-size:11px;font-weight:600;}}
                QPushButton:hover{{background:{t['border']};border-color:{t['accent']};}}
                QPushButton:checked{{background:{t['accent']};color:white;border-color:{t['accent']};}}""")
            btn.setCheckable(True); btn.clicked.connect(lambda _, k=key, b=btn: self.select(k, b))
            grid.addWidget(btn, i//3, i%3); self._btns.append(btn)
        scroll.setWidget(inner); lay.addWidget(scroll)

        row2 = QHBoxLayout(); row2.setSpacing(10)
        cancel = QPushButton("Cancel"); cancel.clicked.connect(self.reject); row2.addWidget(cancel)
        apply = QPushButton("Apply Template"); apply.setObjectName("apply")
        apply.clicked.connect(self.apply_template); row2.addWidget(apply)
        lay.addLayout(row2)

    def select(self, key, btn):
        self.selected_key = key
        for b in self._btns: b.setChecked(False)
        btn.setChecked(True)

    def apply_template(self):
        if not self.selected_key:
            QMessageBox.warning(self,"Warning","Please select a template first."); return
        name = self.name_input.text().strip() or "My Project"
        slug = re.sub(r"[^a-zA-Z0-9-]","-", name.lower()).strip("-")
        tpl  = QUICK_TEMPLATES[self.selected_key]
        self.selected_template = tpl.replace("{name}", name).replace("{slug}", slug)
        self.accept()

# ═══════════════════════════════════════════════════════════════════════════
# SETUP WIZARD
# ═══════════════════════════════════════════════════════════════════════════
class SetupWizard(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        t = T()
        self.setWindowTitle("README Wizard"); self.setFixedSize(600, 520)
        self.setStyleSheet(f"""
            QDialog    {{ background:{t['panel']}; }}
            QLabel     {{ color:{t['text']}; background:transparent; }}
            QLineEdit  {{ background:{t['input_bg']}; color:{t['editor_fg']};
                          border:2px solid {t['border']}; border-radius:8px; padding:12px 14px; font-size:14px; }}
            QLineEdit:focus {{ border-color:{t['accent']}; }}
            QPushButton {{ background:{t['accent']}; color:white; border:none;
                           border-radius:8px; padding:14px; font-weight:600; font-size:15px; }}
            QPushButton:hover {{ background:{t['accent2']}; }}
        """)
        outer_lay = QVBoxLayout(self); outer_lay.setContentsMargins(24,24,24,24); outer_lay.setSpacing(16)
        outer_lay.addWidget(QLabel("Quick README Setup",
            styleSheet=f"font-size:24px;color:{t['text']};margin-bottom:8px;font-weight:600;background:transparent;"))

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        scroll.setStyleSheet("background:transparent;")

        inner = QWidget()
        inner.setStyleSheet("background:transparent;")
        lay = QVBoxLayout(inner); lay.setContentsMargins(0,0,10,0); lay.setSpacing(12)

        self.fields = {}
        for lbl, ph, key in [
            ("Project Name",        "My Awesome Project",           "name"),
            ("Short Description",   "A tool that does X",           "desc"),
            ("Installation Command","pip install myproject",         "install"),
            ("Author Name",         "Your Name",                    "author"),
            ("GitHub Username",     "yourusername",                  "github"),
        ]:
            lay.addWidget(QLabel(lbl, styleSheet=f"color:{t['dim']};font-size:13px;font-weight:600;background:transparent;"))
            inp = QLineEdit(); inp.setPlaceholderText(ph)
            lay.addWidget(inp); self.fields[key] = inp

        scroll.setWidget(inner)
        outer_lay.addWidget(scroll)

        btn = QPushButton("Generate README"); btn.clicked.connect(self.generate)
        outer_lay.addWidget(btn)

    def generate(self):
        v = {k: f.text().strip() or f.placeholderText() for k, f in self.fields.items()}
        slug = re.sub(r"[^a-z0-9-]","-",v["name"].lower()).strip("-")
        self.result = f"""# {v['name']}

[![License](https://img.shields.io/badge/license-MIT-blue?style=flat-square)](#)
[![Python](https://img.shields.io/badge/Python-3.8+-blue?style=flat-square&logo=python)](#)

> {v['desc']}

## Installation
```bash
{v['install']}
```

## Usage
```python
import {slug}
# Your code here
```

## Contributing
Contributions welcome! See [CONTRIBUTING.md](CONTRIBUTING.md)

## License
MIT - [{v['author']}](https://github.com/{v['github']})
"""
        self.accept()


# ═══════════════════════════════════════════════════════════════════════════
# BACKUP MANAGER
# ═══════════════════════════════════════════════════════════════════════════
class BackupManager:
    """Keeps up to 10 timestamped backups per file in ~/.readme_builder_backups/"""
    BACKUP_DIR = os.path.expanduser("~/.readme_builder_backups")
    MAX_BACKUPS = 10

    @classmethod
    def save(cls, workspace_dir: str, content: str):
        os.makedirs(cls.BACKUP_DIR, exist_ok=True)
        slug = workspace_dir.replace(os.sep, "_").replace(":", "").strip("_")[-40:]
        ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
        fname = f"{slug}__{ts}.md"
        try:
            Path(os.path.join(cls.BACKUP_DIR, fname)).write_text(content, encoding="utf-8")
            cls._prune(slug)
        except: pass

    @classmethod
    def _prune(cls, slug: str):
        """Remove oldest backups beyond MAX_BACKUPS for this workspace."""
        files = sorted(
            [f for f in os.listdir(cls.BACKUP_DIR) if f.startswith(slug)],
            reverse=True
        )
        for old in files[cls.MAX_BACKUPS:]:
            try: os.remove(os.path.join(cls.BACKUP_DIR, old))
            except: pass

    @classmethod
    def list_backups(cls, workspace_dir: str):
        if not os.path.exists(cls.BACKUP_DIR): return []
        slug = workspace_dir.replace(os.sep, "_").replace(":", "").strip("_")[-40:]
        files = sorted(
            [f for f in os.listdir(cls.BACKUP_DIR) if f.startswith(slug)],
            reverse=True
        )
        return [os.path.join(cls.BACKUP_DIR, f) for f in files]

    @classmethod
    def read(cls, path: str) -> str:
        try: return Path(path).read_text(encoding="utf-8")
        except: return ""


class BackupDialog(QDialog):
    restore_content = Signal(str)

    def __init__(self, workspace_dir, parent=None):
        super().__init__(parent)
        self.workspace_dir = workspace_dir
        t = T()
        self.setWindowTitle("Backup Versions"); self.setFixedSize(580, 420)
        self.setStyleSheet(f"""
            QDialog    {{ background:{t['panel']}; color:{t['text']}; }}
            QLabel     {{ color:{t['text']}; background:transparent; font-size:13px; }}
            QListWidget {{ background:{t['bg']}; color:{t['editor_fg']};
                           border:1px solid {t['border']}; border-radius:8px;
                           padding:6px; font-size:12px; font-family:monospace; }}
            QListWidget::item:selected {{ background:{t['accent']}; color:white; }}
            QListWidget::item:hover    {{ background:{t['hover']}; }}
            QTextEdit  {{ background:{t['bg']}; color:{t['editor_fg']};
                          border:1px solid {t['border']}; border-radius:8px;
                          padding:8px; font-family:monospace; font-size:11px; }}
            QPushButton {{ background:{t['hover']}; color:{t['text']};
                           border:1px solid {t['border']}; border-radius:8px;
                           padding:8px 14px; font-size:12px; font-weight:600; }}
            QPushButton:hover  {{ background:{t['border']}; border-color:{t['accent']}; }}
            QPushButton#restore {{ background:{t['accent']}; border:none; color:white; }}
            QPushButton#restore:hover {{ background:{t['accent2']}; }}
        """)

        lay = QVBoxLayout(self); lay.setContentsMargins(20,16,20,16); lay.setSpacing(10)
        lay.addWidget(QLabel("Backup Versions",
            styleSheet=f"font-size:16px;font-weight:700;color:{t['text']};background:transparent;"))
        lay.addWidget(QLabel(f"Stored in: {BackupManager.BACKUP_DIR}",
            styleSheet=f"color:{t['dim']};font-size:11px;background:transparent;"))

        splitter = QSplitter(Qt.Horizontal)
        self.backup_list = QListWidget()
        splitter.addWidget(self.backup_list)
        self.preview_box = QTextEdit(); self.preview_box.setReadOnly(True)
        splitter.addWidget(self.preview_box)
        splitter.setSizes([240, 300])
        lay.addWidget(splitter)

        btns = QHBoxLayout(); btns.setSpacing(8)
        refresh_btn = QPushButton("Refresh"); refresh_btn.clicked.connect(self.load_list)
        btns.addWidget(refresh_btn)
        btns.addStretch()
        cancel_btn = QPushButton("Close"); cancel_btn.clicked.connect(self.reject)
        btns.addWidget(cancel_btn)
        restore_btn = QPushButton("Restore Selected"); restore_btn.setObjectName("restore")
        restore_btn.clicked.connect(self._restore); btns.addWidget(restore_btn)
        lay.addLayout(btns)

        self.backup_list.currentRowChanged.connect(self._on_select)
        self.load_list()

    def load_list(self):
        self.backup_list.clear()
        self._paths = BackupManager.list_backups(self.workspace_dir)
        if not self._paths:
            self.backup_list.addItem("  No backups found.")
            return
        for p in self._paths:
            fname = os.path.basename(p)
            # Parse timestamp from filename __YYYYMMDD_HHMMSS.md
            try:
                ts_part = fname.rsplit("__", 1)[-1].replace(".md", "")
                dt = datetime.strptime(ts_part, "%Y%m%d_%H%M%S")
                label = dt.strftime("%Y-%m-%d  %H:%M:%S")
            except:
                label = fname
            self.backup_list.addItem(label)

    def _on_select(self, idx):
        if 0 <= idx < len(self._paths):
            content = BackupManager.read(self._paths[idx])
            self.preview_box.setPlainText(content[:3000] + ("..." if len(content) > 3000 else ""))

    def _restore(self):
        idx = self.backup_list.currentRow()
        if 0 <= idx < len(self._paths):
            content = BackupManager.read(self._paths[idx])
            if content:
                self.restore_content.emit(content); self.accept()


# ═══════════════════════════════════════════════════════════════════════════
# MAIN EDITOR WINDOW
# ═══════════════════════════════════════════════════════════════════════════
# ═══════════════════════════════════════════════════════════════════════════
# PREVIEW BROWSER — QTextBrowser with remote image support
# ═══════════════════════════════════════════════════════════════════════════
class _ImgFetcher(QObject):
    """Downloads a single remote image in a background thread."""
    done = Signal(str, bytes)   # url_str, raw_bytes

    def __init__(self, url_str):
        super().__init__()
        self.url_str = url_str

    def run(self):
        try:
            req = urllib.request.Request(
                self.url_str, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=8) as resp:
                data = resp.read()
            self.done.emit(self.url_str, data)
        except:
            self.done.emit(self.url_str, b"")


class PreviewBrowser(QTextBrowser):
    def __init__(self, workspace_dir="", parent=None):
        super().__init__(parent)
        self.workspace_dir = workspace_dir
        self._img_cache   = {}          # url_str -> QPixmap (ready)
        self._pending     = set()       # urls currently being fetched
        self._threads     = []          # keep threads alive
        self.setOpenExternalLinks(True)
        self.setOpenLinks(True)

    # ── called by Qt whenever the document needs a resource ──────────────
    def loadResource(self, resource_type, url):
        if resource_type != QTextDocument.ImageResource:
            return super().loadResource(resource_type, url)

        url_str = url.toString()

        # 1. Cache hit
        if url_str in self._img_cache:
            return self._img_cache[url_str]

        # 2. Local image — try multiple path resolutions
        candidates = []

        # From QUrl
        local = url.toLocalFile()
        if local:
            candidates.append(local)

        # Normalize slashes both ways
        raw = url_str.replace("\\", "/").replace("%5C", "/").replace("%5c", "/")
        candidates.append(raw)
        candidates.append(raw.replace("/", os.sep))

        # Relative to workspace_dir
        for p in [raw, raw.replace("/", os.sep)]:
            candidates.append(os.path.join(self.workspace_dir, p))

        # Also try with workspace as base for absolute-looking paths
        if raw.startswith("/") or (len(raw) > 1 and raw[1] == ":"):
            candidates.append(raw)

        seen = set()
        for path in candidates:
            if not path:
                continue
            path = os.path.normpath(path)
            if path in seen:
                continue
            seen.add(path)
            if os.path.isfile(path):
                # Read raw bytes so PNG alpha works correctly
                try:
                    with open(path, "rb") as f:
                        data = f.read()
                    img = QImage()
                    img.loadFromData(data)
                    if not img.isNull():
                        px = QPixmap.fromImage(img)
                        self._img_cache[url_str] = px
                        return px
                except:
                    pass

        # 3. Remote image — fetch in background, refresh when done
        if url_str.startswith(("http://", "https://")) and url_str not in self._pending:
            self._pending.add(url_str)
            fetcher = _ImgFetcher(url_str)
            thread  = QThread()
            fetcher.moveToThread(thread)
            fetcher.done.connect(self._on_img_ready)
            thread.started.connect(fetcher.run)
            thread.start()
            self._threads.append((thread, fetcher))

        return QPixmap()

    # ── called from background thread when image bytes arrive ────────────
    def _on_img_ready(self, url_str: str, data: bytes):
        self._pending.discard(url_str)
        if data:
            img = QImage()
            img.loadFromData(data)
            if not img.isNull():
                px = QPixmap.fromImage(img)
                if px.height() < 40:   # badge
                    px = px.scaledToHeight(20, Qt.SmoothTransformation)
                self._img_cache[url_str] = px
                QTimer.singleShot(0, self._refresh_images)

        # Clean up finished threads
        self._threads = [(t, w) for t, w in self._threads if t.isRunning()]

    def _refresh_images(self):
        """Re-render HTML preserving scroll position so images appear."""
        sb  = self.verticalScrollBar()
        pos = sb.value()
        html = self.toHtml()
        self.setHtml(html)
        QTimer.singleShot(0, lambda: sb.setValue(pos))


class ReadmeEditor(QMainWindow):
    def __init__(self, workspace_dir=None, initial_file=None):
        super().__init__()
        self.workspace_dir = os.path.abspath(workspace_dir or os.getcwd())
        self.theme_index = 0
        set_theme_index(0)
        self.setWindowTitle("README Builder")
        self.setFixedSize(1000, 750)
        self.setWindowFlags(Qt.Window)

        register_context_menu()

        icon_path = os.path.join(os.path.dirname(__file__), "icon.png")
        if os.path.exists(icon_path): self.setWindowIcon(QIcon(icon_path))

        central = QWidget(); self.setCentralWidget(central)
        outer_layout = QHBoxLayout(central); outer_layout.setContentsMargins(0,0,0,0); outer_layout.setSpacing(0)
        self.main_splitter = QSplitter(Qt.Horizontal); outer_layout.addWidget(self.main_splitter)
        left_widget = QWidget()
        main_layout = QHBoxLayout(left_widget); main_layout.setContentsMargins(0,0,0,0); main_layout.setSpacing(0)
        self.main_splitter.addWidget(left_widget)

        self.editor = SmartEditor()
        self.emoji_picker = EmojiPicker(self)
        self.emoji_picker.itemClicked.connect(self.insert_emoji)

        self._thread = QThread()
        self._worker = RenderWorker()
        self._worker.moveToThread(self._thread)
        self._worker.finished.connect(self._apply_html)
        self._thread.start()

        self._timer = QTimer(); self._timer.setSingleShot(True)
        self._timer.timeout.connect(self._do_render)
        self.editor.textChanged.connect(lambda: self._timer.start(600))

        # ── File Watcher ─────────────────────────────────────────────────
        from PySide6.QtCore import QFileSystemWatcher
        self._file_watcher = QFileSystemWatcher()
        self._file_watcher.fileChanged.connect(self._on_file_changed_externally)
        self._watcher_suppress = False   # suppress notification after our own save

        self.sidebar = self.create_sidebar()
        main_layout.addWidget(self.sidebar)

        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget); content_layout.setContentsMargins(0,0,0,0); content_layout.setSpacing(0)

        # ── Multi-Tab bar ─────────────────────────────────────────────────
        self._tab_bar_frame = QFrame(); self._tab_bar_frame.setFixedHeight(34)
        tab_bar_lay = QHBoxLayout(self._tab_bar_frame)
        tab_bar_lay.setContentsMargins(6,3,6,0); tab_bar_lay.setSpacing(3)
        self._tab_buttons = []   # list of QToolButton
        self._tabs = []          # list of dict {editor, path, label}
        self._active_tab = 0
        add_tab_btn = QToolButton(); add_tab_btn.setText("+")
        add_tab_btn.setToolTip("New Tab  Ctrl+T")
        add_tab_btn.setStyleSheet(
            f"QToolButton{{background:transparent;color:{T()['accent']};border:none;"
            f"border-radius:6px;padding:4px 10px;font-size:16px;font-weight:700;}}"
            f"QToolButton:hover{{background:{T()['hover']};}}")
        add_tab_btn.clicked.connect(self._new_tab)
        self._tab_bar_lay = tab_bar_lay
        tab_bar_lay.addStretch()
        tab_bar_lay.addWidget(add_tab_btn)
        content_layout.addWidget(self._tab_bar_frame)

        self.setup_topbar(content_layout)

        self.preview = PreviewBrowser(self.workspace_dir)
        self.preview.setSearchPaths([self.workspace_dir])

        self.minimap = Minimap(self.editor)
        self.editor.textChanged.connect(self.minimap.update_from_editor)
        self.editor.verticalScrollBar().valueChanged.connect(self.minimap.update_from_editor)

        # ── Stacked widget holds all tab editors ──────────────────────────
        self._editor_stack = QStackedWidget()
        self._editor_stack.addWidget(self.editor)

        editor_wrapper = QWidget()
        ew_lay = QHBoxLayout(editor_wrapper); ew_lay.setContentsMargins(0,0,0,0); ew_lay.setSpacing(0)
        ew_lay.addWidget(self._editor_stack); ew_lay.addWidget(self.minimap)

        self.splitter = QSplitter(Qt.Horizontal)
        self.splitter.addWidget(editor_wrapper); self.splitter.addWidget(self.preview)
        self.splitter.setSizes([480, 480]); self.splitter.setHandleWidth(1)
        self.splitter.setStyleSheet(f"QSplitter::handle{{background:{T()['border']};}}")
        content_layout.addWidget(self.splitter)

        self.find_bar = FindReplaceBar(self.editor); self.find_bar.hide()
        content_layout.addWidget(self.find_bar)

        self.stats_bar = StatsBar(self.editor)
        content_layout.addWidget(self.stats_bar)

        main_layout.addWidget(content_widget)

        self.right_pane = self.create_right_pane()
        self.main_splitter.addWidget(self.right_pane); self.right_pane.hide()
        self.main_splitter.setSizes([700, 300])

        self.status_lbl = QLabel("", self)
        self.status_lbl.setAlignment(Qt.AlignCenter); self.status_lbl.setFixedSize(140, 40)
        self.status_lbl.setStyleSheet(f"background:{T()['accent']};color:white;border-radius:10px;font-weight:600;font-size:13px;")
        self.status_lbl.hide()

        QShortcut(QKeySequence("Ctrl+S"), self, self.quick_save)
        QShortcut(QKeySequence("Ctrl+F"), self, lambda: (self.find_bar.show(), self.find_bar.find_input.setFocus()))
        QShortcut(QKeySequence("Ctrl+H"), self, lambda: (self.find_bar.show(), self.find_bar.replace_input.setFocus()))
        QShortcut(QKeySequence("Ctrl+M"), self, lambda: self.minimap.setVisible(not self.minimap.isVisible()))
        QShortcut(QKeySequence("Ctrl+Z"), self, self.editor.undo)
        QShortcut(QKeySequence("Ctrl+Y"), self, self.editor.redo)
        QShortcut(QKeySequence("Ctrl+T"), self, self._new_tab)
        QShortcut(QKeySequence("Ctrl+W"), self, self._close_active_tab)
        QShortcut(QKeySequence("Ctrl+Tab"), self, self._next_tab)

        # Initialize first tab record
        self._tabs = [{"editor": self.editor, "path": None, "label": "README.md"}]
        self._active_tab = 0
        self._rebuild_tab_buttons()

        self.preview_ready = True  # always ready with QTextBrowser
        QTimer.singleShot(300, self._update_preview_html)
        self._update_compare_preview_theme()
        self.apply_theme()

        if initial_file and os.path.exists(initial_file):
            try:
                self.editor.setPlainText(Path(initial_file).read_text(encoding='utf-8'))
                add_recent_file(initial_file)
            except: self.load_readme()
        else:
            self.set_default_template()
            QTimer.singleShot(100, self.load_readme)

    # ─────────────────────────────────────────────────────────────────────────
    # MULTI-TAB MANAGEMENT
    # ─────────────────────────────────────────────────────────────────────────
    def _rebuild_tab_buttons(self):
        t = T()
        # Remove old tab buttons (keep stretch and + button at end)
        while self._tab_bar_lay.count() > 2:
            item = self._tab_bar_lay.takeAt(0)
            if item.widget(): item.widget().deleteLater()
        for i, tab in enumerate(self._tabs):
            label = tab.get("label", "README.md")
            btn = QToolButton(); btn.setText(f"  {label}  ")
            btn.setCheckable(True); btn.setChecked(i == self._active_tab)
            is_active = (i == self._active_tab)
            btn.setStyleSheet(f"""
                QToolButton {{
                    background:{''+t['bg'] if is_active else 'transparent'};
                    color:{t['text'] if is_active else t['dim']};
                    border:none; border-bottom:{'2px solid '+t['accent'] if is_active else 'none'};
                    border-radius:0; padding:4px 12px; font-size:12px; font-weight:{'600' if is_active else '400'};
                }}
                QToolButton:hover {{ background:{t['hover']}; color:{t['text']}; }}
            """)
            idx = i
            btn.clicked.connect(lambda _, ix=idx: self._switch_tab(ix))
            self._tab_bar_lay.insertWidget(i, btn)
        self._tab_bar_frame.setStyleSheet(
            f"QFrame{{background:{t['panel']};border-bottom:1px solid {t['border']};}}")

    def _new_tab(self, path=None, content="", label=None):
        new_editor = SmartEditor()
        if content:
            new_editor.setPlainText(content)
            new_editor._last_text = content
        new_editor.textChanged.connect(lambda: self._timer.start(400))
        new_editor.textChanged.connect(self.minimap.update_from_editor)
        new_editor.verticalScrollBar().valueChanged.connect(self.minimap.update_from_editor)
        self._editor_stack.addWidget(new_editor)
        tab_label = label or (os.path.basename(path) if path else f"README_{len(self._tabs)+1}.md")
        self._tabs.append({"editor": new_editor, "path": path, "label": tab_label})
        self._switch_tab(len(self._tabs) - 1)

    def _switch_tab(self, idx):
        if not (0 <= idx < len(self._tabs)): return
        self._active_tab = idx
        tab = self._tabs[idx]
        self.editor = tab["editor"]
        self._editor_stack.setCurrentWidget(self.editor)
        # Reconnect find bar and stats bar
        if hasattr(self, 'find_bar'):
            self.find_bar.editor = self.editor
        if hasattr(self, 'stats_bar'):
            try: self.stats_bar.editor.textChanged.disconnect(self.stats_bar.update_stats)
            except: pass
            self.stats_bar.editor = self.editor
            self.editor.textChanged.connect(self.stats_bar.update_stats)
            self.editor.cursorPositionChanged.connect(self.stats_bar.update_cursor)
            self.stats_bar.update_stats()
        self.minimap.editor = self.editor
        self.minimap.update_from_editor()
        self._rebuild_tab_buttons()
        self._do_render()

    def _close_active_tab(self):
        if len(self._tabs) <= 1: return   # don't close last tab
        idx = self._active_tab
        tab = self._tabs.pop(idx)
        widget = tab["editor"]
        self._editor_stack.removeWidget(widget)
        widget.deleteLater()
        new_idx = min(idx, len(self._tabs) - 1)
        self._active_tab = new_idx
        self._switch_tab(new_idx)

    def _next_tab(self):
        if len(self._tabs) > 1:
            self._switch_tab((self._active_tab + 1) % len(self._tabs))

    def open_file_in_new_tab(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Open File", self.workspace_dir,
            "Markdown Files (*.md *.markdown);;All Files (*.*)")
        if path:
            try:
                content = Path(path).read_text(encoding='utf-8')
                add_recent_file(path)
                self._new_tab(path=path, content=content, label=os.path.basename(path))
                self.show_status("Opened!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Cannot open:\n{e}")

    # ─────────────────────────────────────────────────────────────────────────
    # FILE WATCHER
    # ─────────────────────────────────────────────────────────────────────────
    def _watch_file(self, path: str):
        # Remove all previously watched files
        if self._file_watcher.files():
            self._file_watcher.removePaths(self._file_watcher.files())
        if path and os.path.exists(path):
            self._file_watcher.addPath(path)

    def _on_file_changed_externally(self, path: str):
        if self._watcher_suppress: return
        if not os.path.exists(path): return
        # Find which tab owns this file
        for i, tab in enumerate(self._tabs):
            if tab.get("path") == path:
                msg = QMessageBox(self)
                msg.setWindowTitle("File Changed Externally")
                msg.setText(f"<b>{os.path.basename(path)}</b> was modified outside README Builder.<br>"
                             "Do you want to reload it?")
                msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
                msg.setDefaultButton(QMessageBox.Yes)
                msg.setStyleSheet(
                    f"QMessageBox{{background:{T()['panel']};}}QLabel{{color:{T()['text']};min-width:360px;}}"
                    f"QPushButton{{background:{T()['accent']};color:white;border:none;border-radius:8px;"
                    f"padding:10px 22px;font-weight:600;}}QPushButton:hover{{background:{T()['accent2']};}}")
                if msg.exec() == QMessageBox.Yes:
                    try:
                        content = Path(path).read_text(encoding='utf-8')
                        tab["editor"].setPlainText(content)
                        tab["editor"]._last_text = content
                        self.show_status("Reloaded!")
                    except: pass
                # Re-add path (watcher removes it on change on some OS)
                QTimer.singleShot(500, lambda p=path: self._file_watcher.addPath(p))
                break

    # ─────────────────────────────────────────────────────────────────────────
    # BACKUP
    # ─────────────────────────────────────────────────────────────────────────
    def open_backups(self):
        dlg = BackupDialog(self.workspace_dir, self)
        def _restore(content):
            self.editor._commit_undo()
            self.editor.setPlainText(content)
            self.editor._last_text = content
            QTimer.singleShot(50, self.editor._commit_undo)
            self.show_status("Restored!")
        dlg.restore_content.connect(_restore)
        dlg.exec()

    def _auto_backup(self):
        """Called on every save."""
        BackupManager.save(self.workspace_dir, self.editor.toPlainText())

    # ─────────────────────────────────────────────────────────────────────────
    # TABLE EDITOR
    # ─────────────────────────────────────────────────────────────────────────
    def open_table_editor(self):
        dlg = TableEditorDialog(self)
        dlg.insert_code.connect(self.insert_text)
        dlg.exec()

    # ─────────────────────────────────────────────────────────────────────────
    def create_sidebar(self):
        sidebar = QFrame(); sidebar.setFixedWidth(85)
        sidebar.setStyleSheet("")
        layout = QVBoxLayout(sidebar); layout.setContentsMargins(8,12,8,12); layout.setSpacing(8)
        for label, slot in [
            ("Theme",    self.toggle_theme),
            ("Recent",   self.show_recent_files),
            ("Backups",  self.open_backups),
            ("Images",   self.open_images),
            ("Structure",self.open_structure),
            ("Info",     self.open_info),
            ("Table",    self.open_table_editor),
            ("Draw",     self.open_draw),
            ("Score",    self.show_score),
            ("Import",   self.import_github_readme),
            ("GitHub",   self.open_github),
            ("Translate",self.open_translate),
            ("Snippet",  self.editor.save_custom_snippet),
            ("Emojis",   lambda: self.emoji_picker.show_at_cursor(self.editor)),
            ("TOC",      self.generate_toc),
            ("Quick",    self.open_quick),
            ("Export",   self.open_export),
        ]:
            btn = QToolButton(); btn.setText(label); btn.clicked.connect(slot); layout.addWidget(btn)
        layout.addStretch()
        save_btn = QToolButton(); save_btn.setText("Save")
        save_btn.setStyleSheet(f"QToolButton{{background:{T()['accent']};color:white;border-radius:8px;padding:10px 4px;font-size:12px;font-weight:600;}}")
        save_btn.clicked.connect(self.save_to_folder); layout.addWidget(save_btn)
        return sidebar

    def open_export(self):
        dlg = ExportDialog(self.editor.toPlainText(), self.workspace_dir, self); dlg.exec()
    def open_images(self): dlg = ImagesDialog(self.workspace_dir, self); dlg.insert_code.connect(self.insert_text); dlg.exec()
    def open_structure(self): dlg = StructureDialog(self.workspace_dir, self); dlg.insert_code.connect(self.insert_text); dlg.exec()
    def open_info(self): dlg = InfoDialog(self.workspace_dir, self); dlg.insert_code.connect(self.insert_text); dlg.exec()
    def open_draw(self): DrawingTool(self).exec()

    def show_recent_files(self):
        dlg = RecentFilesDialog(self)
        dlg.file_selected.connect(self._open_recent_file)
        dlg.exec()

    def _open_recent_file(self, path):
        try:
            content = Path(path).read_text(encoding='utf-8')
            self.editor._commit_undo()
            self.editor.setPlainText(content)
            self.editor._last_text = content
            QTimer.singleShot(50, self.editor._commit_undo)
            self.workspace_dir = os.path.dirname(path)
            self.preview.workspace_dir = self.workspace_dir
            self.preview.setSearchPaths([self.workspace_dir])
            add_recent_file(path)
            self.show_status("Loaded!")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Cannot open file:\n{e}")

    def open_github(self):
        dlg = GitHubDialog(self.editor.toPlainText(), self)
        dlg.insert_readme.connect(self._on_github_insert)
        dlg.exec()

    def _on_github_insert(self, content):
        self.editor._commit_undo()
        if "\n" in content and content.strip().startswith("#"):
            self.editor.setPlainText(content)
            self.editor._last_text = content
        else:
            self.editor.insertPlainText("\n\n" + content)
        QTimer.singleShot(50, self.editor._commit_undo)

    def open_translate(self):
        dlg = TranslateDialog(self.editor.toPlainText(), self)
        def _apply_translation(text):
            self.editor._commit_undo()
            self.editor.setPlainText(text)
            self.editor._last_text = text
            QTimer.singleShot(50, self.editor._commit_undo)
        dlg.translated.connect(_apply_translation)
        dlg.exec()

    def show_score(self):
        score, feedback = calculate_readme_score(self.editor.toPlainText())
        status = "Professional" if score >= 80 else "Needs Improvement"
        msg = f"<div style='text-align:center;padding:20px;'><h2 style='color:#8b5cf6;'>README Score</h2><h1 style='color:#4ade80;font-size:52px;margin:24px 0'>{score}/100</h1><p style='color:#9ca3af;font-size:17px'>{status}</p>"
        if feedback:
            msg += f"<div style='text-align:left;background:{T()['bg']};padding:20px;border-radius:12px'><h3 style='color:#fff'>Suggestions:</h3><ul style='color:#c9d1d9;line-height:1.8'>"
            for item in feedback: msg += f"<li style='margin:10px 0'>{item}</li>"
            msg += "</ul></div>"
        else: msg += "<p style='color:#4ade80;font-size:18px'>Perfect README!</p>"
        msg += "</div>"
        dlg = QMessageBox(self); dlg.setWindowTitle("README Score"); dlg.setTextFormat(Qt.RichText); dlg.setText(msg)
        dlg.setStyleSheet(f"QMessageBox{{background:{T()['panel']};}}QLabel{{color:{T()['text']};min-width:500px;}}QPushButton{{background:{T()['accent']};color:white;border:none;border-radius:8px;padding:12px 28px;font-weight:600;}}QPushButton:hover{{background:{T()['accent2']};}}")
        dlg.exec()

    def open_quick(self):
        dlg = QuickProjectDialog(self); dlg.move(self.mapToGlobal(QPoint(self.width()//2-325, 80)))
        if dlg.exec() == QDialog.Accepted and dlg.selected_template:
            self.editor._commit_undo()          # snapshot before replacing
            self.editor.setPlainText(dlg.selected_template)
            self.editor._last_text = dlg.selected_template
            QTimer.singleShot(50, self.editor._commit_undo)

    def insert_text(self, text): self.editor.insertPlainText(text)

    def setup_topbar(self, parent_layout):
        self.top_bar = QFrame(); self.top_bar.setFixedHeight(52)
        self.top_bar.setStyleSheet("")
        lay = QHBoxLayout(self.top_bar); lay.setContentsMargins(16,0,16,0); lay.setSpacing(10)
        title = QLabel(tr("app_title")); title.setStyleSheet(f"color:{T()['text']};font-weight:600;font-size:15px;background:transparent;"); lay.addWidget(title)
        lay.addStretch()
        for label, obj_name, slot in [
            (" Wizard","wizard",self.run_wizard),
            (" Open Tab","",self.open_file_in_new_tab),
            (" Find  Ctrl+F","",lambda:(self.find_bar.show(),self.find_bar.find_input.setFocus())),
            (" Map  Ctrl+M","",lambda:self.minimap.setVisible(not self.minimap.isVisible())),
            (" Compare","",self.toggle_right_pane),
        ]:
            btn = QPushButton(label)
            if obj_name: btn.setObjectName(obj_name)
            btn.clicked.connect(slot); lay.addWidget(btn)
        parent_layout.addWidget(self.top_bar)

    def run_wizard(self):
        wiz = SetupWizard(self)
        if wiz.exec() == QDialog.Accepted:
            self.editor._commit_undo()          # snapshot before replacing
            self.editor.setPlainText(wiz.result)
            self.editor._last_text = wiz.result
            QTimer.singleShot(50, self.editor._commit_undo)

    def closeEvent(self, e): self._thread.quit(); self._thread.wait(1000); super().closeEvent(e)
    def _do_render(self): self._worker.request(self.editor.toPlainText()); self._worker.process()

    def _update_compare_preview_theme(self):
        if not hasattr(self,'compare_preview') or not hasattr(self,'compare_code'): return
        md_text = self.compare_code.toPlainText()
        if not md_text: return
        html_content = markdown2.markdown(md_text, extras=["fenced-code-blocks","tables"])
        t   = THEMES[self.theme_index]
        css = _preview_css(t)
        self.compare_preview.setHtml(
            f"<!DOCTYPE html><html><head><meta charset='UTF-8'>{css}</head>"
            f"<body><div id='content'>{html_content}</div></body></html>"
        )

    def toggle_theme(self):
        self.theme_index = (self.theme_index + 1) % len(THEMES)
        set_theme_index(self.theme_index)
        self.editor.highlighter.rehighlight()
        self.stats_bar.apply_theme()
        self.find_bar.apply_theme()
        self._rebuild_tab_buttons()
        self.apply_theme(); self._update_preview_html(); self._update_compare_preview_theme()

    def _update_preview_html(self):
        t   = THEMES[self.theme_index]
        css = _preview_css(t)
        html = self._worker._last_html if hasattr(self._worker, '_last_html') else ""
        self.preview.setHtml(
            f"<!DOCTYPE html><html><head><meta charset='UTF-8'>{css}</head>"
            f"<body><div id='content'>{html}</div></body></html>"
        )

    def apply_theme(self):
        t   = THEMES[self.theme_index]
        bg          = t["bg"]
        panel_bg    = t["panel"]
        border_col  = t["border"]
        btn_hover   = t["hover"]
        text        = t["text"]
        text_dim    = t["dim"]
        text_editor = t["editor_fg"]
        input_bg    = t["input_bg"]
        scrollbar   = t["scroll"]
        accent      = t["accent"]
        find_bg     = t["find_bg"]

        self.setStyleSheet(f"""
            QMainWindow {{ background:{bg}; }}
            QWidget     {{ background:{bg}; color:{text}; }}
            QTextEdit   {{
                background:{bg}; color:{text_editor}; border:none;
                font-size:15px; padding:20px;
                font-family:ui-monospace,'SF Mono',Consolas,monospace;
                line-height:1.7;
            }}
            QScrollBar:vertical   {{ border:none; background:{bg}; width:8px; margin:0; }}
            QScrollBar::handle:vertical       {{ background:{scrollbar}; min-height:30px; border-radius:4px; }}
            QScrollBar::handle:vertical:hover {{ background:{border_col}; }}
            QScrollBar:horizontal {{ border:none; background:{bg}; height:8px; margin:0; }}
            QScrollBar::handle:horizontal       {{ background:{scrollbar}; min-width:30px; border-radius:4px; }}
            QScrollBar::handle:horizontal:hover {{ background:{border_col}; }}
        """)

        if hasattr(self, 'sidebar'):
            self.sidebar.setStyleSheet(f"""
                QFrame      {{ background:{panel_bg}; border-right:1px solid {border_col}; }}
                QToolButton {{
                    background:transparent; color:{text_dim};
                    border:none; border-radius:10px;
                    padding:10px 4px; font-size:11px; font-weight:600;
                }}
                QToolButton:hover   {{ background:{btn_hover}; color:{text}; }}
                QToolButton:checked {{ background:{accent}; color:white; }}
            """)

        if hasattr(self, 'top_bar'):
            self.top_bar.setStyleSheet(f"""
                QFrame      {{ background:{panel_bg}; border-bottom:1px solid {border_col}; }}
                QLabel      {{ color:{text}; background:transparent; }}
                QPushButton {{
                    background:{btn_hover}; color:{text};
                    border:1px solid {border_col}; border-radius:8px;
                    padding:8px 14px; font-size:12px; font-weight:600;
                }}
                QPushButton:hover  {{ background:{border_col}; border-color:{accent}; }}
                QPushButton#wizard {{ background:{accent}; border:none; color:white; }}
                QPushButton#wizard:hover {{ background:#7c3aed; }}
            """)

        if hasattr(self, 'find_bar'):
            self.find_bar.setStyleSheet(f"""
                QFrame      {{ background:{find_bg}; border-top:1px solid {border_col}; }}
                QLineEdit   {{
                    background:{input_bg}; color:{text_editor};
                    border:1px solid {border_col}; border-radius:6px;
                    padding:6px 10px; font-size:12px;
                }}
                QLineEdit:focus {{ border-color:{accent}; }}
                QPushButton {{
                    background:{btn_hover}; color:{text};
                    border:1px solid {border_col}; border-radius:6px;
                    padding:6px 10px; font-size:11px; font-weight:600;
                }}
                QPushButton:hover {{ background:{border_col}; border-color:{accent}; }}
                QLabel    {{ color:{text_dim}; font-size:11px; background:transparent; }}
                QCheckBox {{ color:{text_dim}; font-size:11px; background:transparent; }}
            """)

        if hasattr(self, 'stats_bar'):
            self.stats_bar.setStyleSheet(f"""
                QFrame {{ background:{panel_bg}; border-top:1px solid {border_col}; }}
                QLabel {{ font-size:11px; font-family:'Consolas','SF Mono',monospace;
                          padding:0 4px; background:transparent; color:{text_dim}; }}
            """)
            self.stats_bar.update_stats()
            self.stats_bar.update_cursor()

        if hasattr(self, 'preview'):
            self.preview.setStyleSheet(f"background:{bg};")

        if hasattr(self, 'compare_tabs'):
            self.compare_tabs.setStyleSheet(f"""
                QTabWidget::pane {{ border-left:1px solid {border_col}; background:{bg}; }}
                QTabBar::tab {{
                    background:{panel_bg}; color:{text_dim};
                    padding:8px 16px; border:1px solid {border_col};
                    border-bottom:none;
                    border-top-left-radius:4px; border-top-right-radius:4px;
                }}
                QTabBar::tab:selected  {{ background:{bg}; color:{text}; border-bottom:1px solid {bg}; }}
                QTabBar::tab:hover     {{ background:{btn_hover}; color:{text}; }}
            """)
            self.compare_code.setStyleSheet(
                f"background:{bg}; color:{text_editor}; border:none; font-family:monospace; padding:10px;"
            )
            if hasattr(self, 'compare_url_input'):
                self.compare_url_input.setStyleSheet(f"""
                    QLineEdit {{
                        background:{input_bg}; color:{text};
                        border:1px solid {border_col}; border-radius:6px; padding:6px;
                    }}
                    QLineEdit:focus {{ border-color:{accent}; }}
                """)

        if hasattr(self, 'minimap'):
            self.minimap.setStyleSheet(
                f"background:{t['line_num_bg']}; border-left:1px solid {border_col};"
            )

        if hasattr(self, 'compare_top_bar'):
            self._style_right_pane(t)

    def _apply_html(self, html):
        self._worker._last_html = html
        t   = THEMES[self.theme_index]
        css = _preview_css(t)
        # Fix Windows backslashes in img src so QTextBrowser can resolve them
        import re as _re
        html = _re.sub(
            r'(<img[^>]+src=["\'])([^"\']+)(["\'])',
            lambda m: m.group(1) + m.group(2).replace("\\", "/") + m.group(3),
            html
        )
        self.preview.setHtml(
            f"<!DOCTYPE html><html><head><meta charset='UTF-8'>{css}</head>"
            f"<body><div id='content'>{html}</div></body></html>"
        )
        self._update_compare_preview_theme()

    def toggle_right_pane(self):
        if self.right_pane.isHidden():
            self.right_pane.show(); self.setFixedSize(1400, 750); self.main_splitter.setSizes([1000, 400])
        else:
            self.right_pane.hide(); self.setFixedSize(1000, 750); self.main_splitter.setSizes([1000, 0])

    def create_right_pane(self):
        t   = THEMES[self.theme_index]
        pane = QWidget(); pane.setMinimumWidth(300)
        lay = QVBoxLayout(pane); lay.setContentsMargins(0,0,0,0); lay.setSpacing(0)

        self.compare_top_bar = QFrame(); self.compare_top_bar.setFixedHeight(52)
        tb_lay = QHBoxLayout(self.compare_top_bar); tb_lay.setContentsMargins(10,0,10,0)
        self.compare_url_input = QLineEdit()
        self.compare_url_input.setPlaceholderText("https://github.com/user/repo")
        tb_lay.addWidget(self.compare_url_input)
        fetch_btn = QPushButton("Load"); fetch_btn.clicked.connect(self.fetch_compare_repo)
        tb_lay.addWidget(fetch_btn)
        lay.addWidget(self.compare_top_bar)

        self.compare_tabs = QTabWidget()
        self.compare_code = QTextEdit(); self.compare_code.setReadOnly(True)
        self.compare_tabs.addTab(self.compare_code, "Code")
        self.compare_preview = QTextBrowser()
        self.compare_preview.setOpenExternalLinks(True)
        self.compare_tabs.addTab(self.compare_preview, "Preview")
        lay.addWidget(self.compare_tabs)

        self._style_right_pane(t)
        return pane

    def _style_right_pane(self, t):
        bg         = t["bg"];       panel_bg  = t["panel"]
        border_col = t["border"];   btn_hover = t["hover"]
        text       = t["text"];     text_dim  = t["dim"]
        input_bg   = t["input_bg"]; accent    = t["accent"]

        if hasattr(self, 'compare_top_bar'):
            self.compare_top_bar.setStyleSheet(f"""
                QFrame      {{ background:{panel_bg}; border-bottom:1px solid {border_col};
                               border-left:1px solid {border_col}; }}
                QLineEdit   {{
                    background:{input_bg}; color:{text};
                    border:1px solid {border_col}; border-radius:6px; padding:6px;
                }}
                QLineEdit:focus {{ border-color:{accent}; }}
                QPushButton {{
                    background:{btn_hover}; color:{text};
                    border:1px solid {border_col}; border-radius:6px;
                    padding:6px 12px; font-weight:bold;
                }}
                QPushButton:hover {{ background:{border_col}; border-color:{accent}; }}
            """)

    def fetch_compare_repo(self):
        url = self.compare_url_input.text().strip()
        if not url: return
        raw = url
        if "github.com" in raw and "raw.githubusercontent" not in raw:
            raw = raw.replace("github.com","raw.githubusercontent.com")
            raw = raw.replace("/blob/","/") if "/blob/" in raw else raw.rstrip("/")+"/main/README.md"
        try:
            try: resp = urllib.request.urlopen(urllib.request.Request(raw))
            except urllib.error.HTTPError as e:
                if e.code == 404 and "/main/" in raw: resp = urllib.request.urlopen(urllib.request.Request(raw.replace("/main/","/master/")))
                else: raise
            md_text = resp.read().decode("utf-8")
            self.compare_code.setPlainText(md_text)
            html_content = markdown2.markdown(md_text, extras=["fenced-code-blocks","tables"])
            t = THEMES[self.theme_index]
            css = _preview_css(t)
            self.compare_preview.setHtml(f"<!DOCTYPE html><html><head><meta charset='UTF-8'>{css}</head><body><div id='content'>{html_content}</div></body></html>")
        except Exception as e: self.compare_code.setPlainText(f"Error loading README:\n{str(e)}")

    def save_to_folder(self):
        if self.perform_save(): self.close()

    def quick_save(self):
        if self.perform_save():
            fp = os.path.join(self.workspace_dir, "README.md")
            add_recent_file(fp)
            self._auto_backup()
            # Update watcher
            tab = self._tabs[self._active_tab] if self._tabs else {}
            if tab.get("path"):
                self._watcher_suppress = True
                QTimer.singleShot(800, lambda: setattr(self, '_watcher_suppress', False))
                self._watch_file(tab["path"])
            self.show_status(tr("save"))

    def perform_save(self):
        fp = os.path.join(self.workspace_dir, "README.md")
        try:
            with open(fp,"w",encoding="utf-8") as f: f.write(self.editor.toPlainText())
            return True
        except Exception as e: QMessageBox.critical(self,"Error",f"Failed to save:\n{e}"); return False

    def show_status(self, msg):
        self.status_lbl.setText(msg)
        self.status_lbl.move(self.width()//2-70, self.height()-60)
        self.status_lbl.show(); QTimer.singleShot(1600, self.status_lbl.hide)

    def import_github_readme(self):
        url, ok = ModernInputDialog.get_text(self,"Import from GitHub","GitHub Repo URL:","https://github.com/user/repo")
        if not (ok and url): return
        raw = url.strip()
        if "github.com" in raw and "raw.githubusercontent" not in raw:
            raw = raw.replace("github.com","raw.githubusercontent.com")
            raw = raw.replace("/blob/","/") if "/blob/" in raw else raw.rstrip("/")+"/main/README.md"
        try:
            try: resp = urllib.request.urlopen(urllib.request.Request(raw))
            except urllib.error.HTTPError as e:
                if e.code == 404 and "/main/" in raw: resp = urllib.request.urlopen(urllib.request.Request(raw.replace("/main/","/master/")))
                else: raise
            self.editor.setPlainText(resp.read().decode("utf-8"))
        except Exception as e: QMessageBox.warning(self,"Error",f"Error\n{e}")

    def insert_emoji(self, item):
        self.editor.insertPlainText(item.text()); self.emoji_picker.hide(); self.editor.setFocus()

    def generate_toc(self):
        lines = self.editor.toPlainText().split("\n")
        toc = ["## Table of Contents"]; found = False
        for line in lines:
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                title = line.strip("#").strip()
                anchor = title.lower().replace(" ","-").replace(".","").replace(":","")
                toc.append("  "*(level-1) + f"- [{title}](#{anchor})")
                found = True
        if found: self.editor.setPlainText("\n".join(toc) + "\n\n\n\n" + self.editor.toPlainText())
        else: QMessageBox.information(self,"TOC","No headers found.")

    def set_default_template(self):
        self.editor.setPlainText("""# Project Name

## Description
Write a clear description here.

## Installation
```bash
pip install -r requirements.txt
```

## Features
- [x] Feature one
- [x] Fast & Minimalistic

## Screenshots
![Preview](example/1.png)
""")

    def load_readme(self):
        path = Path(self.workspace_dir) / "README.md"
        if path.exists():
            try:
                self.editor.setPlainText(path.read_text(encoding='utf-8'))
                add_recent_file(str(path))
            except: pass


# ══════════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════
def register_context_menu():
    if sys.platform == 'win32':
        try:
            import winreg
            exe = sys.executable
            is_frozen = getattr(sys, 'frozen', False)
            if not is_frozen:
                cmd_prefix = f'"{exe}" "{os.path.abspath(__file__)}"'
            else:
                cmd_prefix = f'"{exe}"'

            for key_path, val, cmd_val in [
                (r"Software\Classes\*\shell\Open with README Builder", "Edit with README Builder", f'{cmd_prefix} "%1"'),
                (r"Software\Classes\Directory\Background\shell\Create README", "Create README here", f'{cmd_prefix} "%V"'),
            ]:
                with winreg.CreateKey(winreg.HKEY_CURRENT_USER, key_path) as k:
                    winreg.SetValue(k, "", winreg.REG_SZ, val)
                    winreg.SetValueEx(k, "Icon", 0, winreg.REG_SZ, exe)
                with winreg.CreateKey(winreg.HKEY_CURRENT_USER, key_path + r"\command") as k:
                    winreg.SetValue(k, "", winreg.REG_SZ, cmd_val)
        except Exception as e:
            print(f"Registry error: {e}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    app.setStyleSheet(f"""
        QWidget{{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;}}
        QMessageBox{{background:{T()['panel']};}}
        QMessageBox QLabel{{color:{T()['text']};}}
        QToolTip{{background:{T()['panel']};color:{T()['text']};border:1px solid {T()['border']};border-radius:6px;padding:4px 8px;font-size:12px;}}
    """)
    icon_path = os.path.join(os.path.dirname(__file__), "icon.png")
    if os.path.exists(icon_path): app.setWindowIcon(QIcon(icon_path))

    workspace_dir = os.getcwd()
    initial_file = None

    if len(sys.argv) > 1:
        arg = sys.argv[1]
        if arg == "--register":
            register_context_menu()
            print("Context menu registered!"); sys.exit(0)

        path = os.path.abspath(arg)
        if os.path.isdir(path):
            workspace_dir = path
        elif os.path.isfile(path):
            workspace_dir = os.path.dirname(path)
            initial_file = path

    window = ReadmeEditor(workspace_dir, initial_file)
    window.show()
    sys.exit(app.exec())